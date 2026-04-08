"""
生成论文格式 Word 回归结果表（三线表）
输出：results/回归结果表.docx
"""
import os
import sys
import warnings
import numpy as np

warnings.filterwarnings("ignore")

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def _set_run_font(run, font_name="Times New Roman", east_asia="宋体"):
    """安全设置 run 的中西文字体"""
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), east_asia)

sys.path.insert(0, os.path.join(os.getcwd(), "scripts"))
from regression_analysis import (
    build_analysis_dataset,
    run_regressions,
    _fit_expectation_salary_model,
    _build_fe_matrix,
    _fit_with_cluster_se,
)


# ============================================================
# 辅助函数
# ============================================================

def format_coef(coef, pval):
    """格式化系数，附加显著性星号"""
    stars = ""
    if pval < 0.01:
        stars = "***"
    elif pval < 0.05:
        stars = "**"
    elif pval < 0.1:
        stars = "*"
    return f"{coef:.4f}{stars}"


def format_tval(tval):
    return f"({tval:.4f})"


def format_pval(pval):
    if pval is None or (isinstance(pval, float) and np.isnan(pval)):
        return "nan"
    if pval < 0.001:
        return "<0.001"
    return f"{pval:.4f}"


def _get_tstat(model, var_name):
    if hasattr(model, "tstats") and var_name in model.tstats:
        return model.tstats[var_name]
    return model.tvalues[var_name]


def _get_std_error(model, var_name):
    if hasattr(model, "std_errors") and var_name in model.std_errors:
        return model.std_errors[var_name]
    return model.bse[var_name]


def _get_fstat(model):
    stat = getattr(model, "f_statistic_robust", None)
    if stat is not None and getattr(stat, "stat", None) is not None:
        return float(stat.stat)
    stat = getattr(model, "f_statistic", None)
    if stat is not None and getattr(stat, "stat", None) is not None:
        return float(stat.stat)
    return np.nan


def fit_model(df, dep_var, core_vars, df_pre=None):
    """回归拟合"""
    if df_pre is not None:
        sub = df_pre
    else:
        needed = core_vars + [dep_var, "IndustrySector", "Year"]
        sub = df.dropna(subset=needed).copy()
    if len(sub) == 0:
        return None, 0
    sub = sub.set_index(["Symbol", "Year"], drop=False)
    x, _, _ = _build_fe_matrix(sub, core_vars)
    y = sub[dep_var]
    model = _fit_with_cluster_se(y, x, sub["Symbol"])
    return model, len(sub)


# ============================================================
# Word 三线表样式
# ============================================================

def set_cell_text(cell, text, bold=False, font_size=10.5, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格文本和格式"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    _set_run_font(run, "Times New Roman", "宋体")
    run.bold = bold
    # 段落间距设为0
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = Pt(14)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val["val"])
        element.set(qn('w:sz'), val["sz"])
        element.set(qn('w:space'), "0")
        element.set(qn('w:color'), val.get("color", "000000"))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def apply_three_line_style(table, header_rows=2):
    """对表格应用三线表样式：顶部粗线、表头下粗线、底部粗线"""
    # 先清除所有默认边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}>')
    # 清除表格边框
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    n_rows = len(table.rows)
    n_cols = len(table.columns)
    thick = {"val": "single", "sz": "12"}  # 1.5pt
    thin = {"val": "single", "sz": "6"}    # 0.75pt

    # 顶部粗线（第 0 行的上边框）
    for col in range(n_cols):
        set_cell_border(table.cell(0, col), top=thick)

    # 表头下划线（第 header_rows-1 行的下边框）
    for col in range(n_cols):
        set_cell_border(table.cell(header_rows - 1, col), bottom=thin)

    # 底部粗线（最后一行的下边框）
    for col in range(n_cols):
        set_cell_border(table.cell(n_rows - 1, col), bottom=thick)


def add_table_title(doc, title):
    """添加表格标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(12)
    _set_run_font(run, "Times New Roman", "黑体")
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_table_note(doc, note):
    """添加表格注释"""
    p = doc.add_paragraph()
    run = p.add_run(note)
    run.font.size = Pt(9)
    _set_run_font(run, "Times New Roman", "宋体")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)


# ============================================================
# 各表格构建
# ============================================================

def build_desc_table_docx(doc, df):
    """表1 描述性统计"""
    add_table_title(doc, "表1  主要变量描述性统计")

    desc_vars_map = {
        "Top3Salary": "高管前三名薪酬总额",
        "SubsidyAmount": "政府补助",
        "lnSubsidy": "财政补贴强度(lnSubsidy)",
        "lnCEOpay": "薪酬对数(lnCEOpay)",
        "lnSale": "企业规模(lnSale)",
        "IA": "无形资产占比(IA)",
        "Overpay": "超额薪酬(Overpay)",
        "Power": "管理层权力(Power, FA)",
        "Roa": "资产收益率(Roa)",
        "Lever": "财务杠杆(Lever)",
        "Top1": "第一大股东持股比例(Top1)",
        "Zone": "地区(Zone, 中西部=1)",
    }

    headers = ["变量", "N", "均值", "中位数", "标准差", "最小值", "最大值"]
    rows_data = []
    for var, label in desc_vars_map.items():
        if var in df.columns:
            s = df[var].dropna()
            def fmt(v):
                return f"{v:.4f}" if abs(v) < 1e6 else f"{v:.2e}"
            rows_data.append([label, str(int(s.count())), fmt(s.mean()), fmt(s.median()), fmt(s.std()), fmt(s.min()), fmt(s.max())])

    table = doc.add_table(rows=1 + len(rows_data), cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)

    # 数据行
    for r, row in enumerate(rows_data):
        for c, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(table.cell(r + 1, c), val, alignment=align)

    apply_three_line_style(table, header_rows=1)
    add_table_note(doc, "注：薪酬和政府补助单位为元。Overpay 直接使用模型1残差，不再进行二次缩尾；其余连续变量按 1%/99% 分位数缩尾处理。Zone 为虚拟变量，其中中西部=1、东部=0。")


def build_first_stage_table_docx(doc, df):
    """表2 第一阶段期望薪酬模型。"""
    fit_result = _fit_expectation_salary_model(df)
    model = fit_result["model"]
    n = len(fit_result["df_model1"])

    explanations = {
        "lnSale": "企业规模越大，期望薪酬越高",
        "Roa": "盈利能力越强，期望薪酬越高",
        "IA": "无形资产占比较高时，期望薪酬相对较低",
        "Zone": "在中西部=1的定义下，负系数表示中西部期望薪酬低于东部",
    }
    display_vars = ["lnSale", "Roa", "IA", "Zone"]

    add_table_title(doc, "表2  第一阶段期望薪酬模型估计结果")
    table = doc.add_table(rows=1 + len(display_vars) + 4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["变量", "系数", "t值", "说明"]
    for i, header in enumerate(headers):
        set_cell_text(table.cell(0, i), header, bold=True)

    row_idx = 1
    for var in display_vars:
        set_cell_text(table.cell(row_idx, 0), var, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.cell(row_idx, 1), format_coef(model.params[var], model.pvalues[var]) if model else "—")
        set_cell_text(table.cell(row_idx, 2), format_tval(_get_tstat(model, var)) if model else "")
        set_cell_text(table.cell(row_idx, 3), explanations[var], alignment=WD_ALIGN_PARAGRAPH.LEFT)
        row_idx += 1

    tail_rows = [
        ("Industry FE", "控制", "", "17 个行业虚拟变量"),
        ("Year FE", "控制", "", "21 个年份虚拟变量"),
        ("N", str(n), "", "—"),
        ("R²", f"{model.rsquared:.4f}" if model else "—", "", "—"),
    ]
    for label, value, extra, note in tail_rows:
        set_cell_text(table.cell(row_idx, 0), label, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.cell(row_idx, 1), value)
        set_cell_text(table.cell(row_idx, 2), extra)
        set_cell_text(table.cell(row_idx, 3), note, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        row_idx += 1

    apply_three_line_style(table, header_rows=1)
    add_table_note(doc, "注：括号中为 t 值，* p<0.1, ** p<0.05, *** p<0.01。模型1使用 OLS 并控制行业与年份虚拟变量，其残差直接记为超额薪酬 Overpay。Zone 为虚拟变量，其中中西部=1、东部=0。")


def _add_regression_table(doc, title, col_headers, sub_headers, display_vars, models_data, note, extra_rows=None):
    """
    通用回归结果表
    models_data: [(model_object, n_sample), ...]
    extra_rows: [(label, [values...]), ...] 额外行（如 Controls, FE）
    """
    add_table_title(doc, title)

    n_models = len(models_data)
    n_cols = 1 + n_models  # 变量列 + 各模型列

    # 计算行数：表头2行 + 每个变量2行(系数+t值) + 分隔行 + extra + N + R²
    n_extra = len(extra_rows) if extra_rows else 0
    n_data_rows = len(display_vars) * 2
    n_total = 2 + n_data_rows + n_extra + 2  # header(2) + vars + extra + N + R²

    table = doc.add_table(rows=n_total, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头行1
    set_cell_text(table.cell(0, 0), "变量", bold=True)
    for i, h in enumerate(col_headers):
        set_cell_text(table.cell(0, i + 1), h, bold=True)

    # 表头行2（因变量）
    set_cell_text(table.cell(1, 0), "", bold=True)
    for i, sh in enumerate(sub_headers):
        set_cell_text(table.cell(1, i + 1), sh, bold=True)

    # 变量行
    row_idx = 2
    for var in display_vars:
        # 系数行
        set_cell_text(table.cell(row_idx, 0), var, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        for m_idx, (model, _) in enumerate(models_data):
            if model is not None and var in model.params:
                text = format_coef(model.params[var], model.pvalues[var])
            else:
                text = ""
            set_cell_text(table.cell(row_idx, m_idx + 1), text)
        row_idx += 1

        # t值行
        set_cell_text(table.cell(row_idx, 0), "")
        for m_idx, (model, _) in enumerate(models_data):
            if model is not None and var in model.params:
                text = format_tval(_get_tstat(model, var))
            else:
                text = ""
            set_cell_text(table.cell(row_idx, m_idx + 1), text)
        row_idx += 1

    # 额外行（Controls, FE 等）
    if extra_rows:
        for label, values in extra_rows:
            set_cell_text(table.cell(row_idx, 0), label, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            for i, v in enumerate(values):
                set_cell_text(table.cell(row_idx, i + 1), v)
            row_idx += 1

    # N 行
    set_cell_text(table.cell(row_idx, 0), "N", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    for i, (_, n) in enumerate(models_data):
        set_cell_text(table.cell(row_idx, i + 1), str(n))
    row_idx += 1

    # R² 行
    set_cell_text(table.cell(row_idx, 0), "R²", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    for i, (model, _) in enumerate(models_data):
        r2 = f"{model.rsquared:.4f}" if model else "—"
        set_cell_text(table.cell(row_idx, i + 1), r2)

    apply_three_line_style(table, header_rows=2)
    add_table_note(doc, note)


def build_main_table_docx(doc, main_results):
    """表3 主回归（模型2）"""
    model2 = main_results["model2"]
    main_fit = main_results["main_result"]
    display_vars = ["lnSubsidy_l1", "Roa", "Lever", "Top1", "Zone"]

    _add_regression_table(
        doc,
        "表3  主回归结果（模型2，公司与年份固定效应）",
        ["模型2"],
        ["Overpay"],
        display_vars,
        [(model2, main_fit["sample_size"])],
        "注：括号中为 t 值，* p<0.1, ** p<0.05, *** p<0.01。F 统计量报告公司层面聚类稳健标准误下的联合显著性检验。标准误为公司层面聚类稳健标准误。",
        extra_rows=[
            ("Firm FE", ["控制"]),
            ("Year FE", ["控制"]),
            ("F统计量", [f"{_get_fstat(model2):.2f}"]),
        ],
    )


def build_mediation_table_docx(doc, df):
    """表5 全样本中介效应检验。"""
    results = df if isinstance(df, dict) else run_regressions(df)
    summary = results["summary"]
    model3 = results["model3"]
    model4 = results["model4"]
    model5 = results["model5"]
    add_table_title(doc, "表5  管理层权力的中介效应检验结果（模型3至模型5，FA口径）")

    rows = [
        ("模型3 总效应 c：lnSubsidy_l1 → Overpay", format_coef(summary["coef_c"], summary["p_c"]), f"{_get_std_error(model3, 'lnSubsidy_l1'):.4f}", "—"),
        ("模型4 路径 a：lnSubsidy_l1 → Power（FA）", format_coef(summary["coef_a"], summary["p_a"]), f"{_get_std_error(model4, 'lnSubsidy_l1'):.4f}", "—"),
        ("模型5 路径 b：Power（FA） → Overpay", format_coef(summary["coef_b"], summary["p_b"]), f"{_get_std_error(model5, 'Power'):.4f}", "—"),
        ("模型5 直接效应 c'：lnSubsidy_l1 → Overpay", format_coef(summary["coef_c_prime"], summary["p_c_prime"]), f"{_get_std_error(model5, 'lnSubsidy_l1'):.4f}", "—"),
        ("间接效应 a×b", f"{summary['indirect_effect']:.6f}", "—", f"[{summary['bootstrap_ci_lower']:.6f}, {summary['bootstrap_ci_upper']:.6f}]"),
        ("Sobel p 值", f"{summary['sobel_p']:.4f}", "—", "—"),
        ("中介效应占比（a×b / c）", f"{summary['mediation_ratio_pct']:.2f}%", "—", "—"),
    ]

    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["路径", "系数", "标准误（cluster）", "Bootstrap 95%CI"]
    for i, header in enumerate(headers):
        set_cell_text(table.cell(0, i), header, bold=True)
    for row_idx, row in enumerate(rows, start=1):
        set_cell_text(table.cell(row_idx, 0), row[0], alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.cell(row_idx, 1), row[1])
        set_cell_text(table.cell(row_idx, 2), row[2])
        set_cell_text(table.cell(row_idx, 3), row[3])

    apply_three_line_style(table, header_rows=1)
    add_table_note(
        doc,
        "注：***、**、* 分别表示在1%、5%、10%水平上显著。模型3报告总效应，模型4报告路径a，模型5同时报告路径b与直接效应。Bootstrap 为公司层面 cluster bootstrap，300 次重抽样。所有回归均控制 Roa、Lever、Top1、Zone 以及公司和年份固定效应。"
    )


def build_causal_table_docx(doc, main_results):
    """表4 FE-2SLS 因果识别结果。"""
    iv_result = main_results["iv_result"]
    first_stage = iv_result["first_stage"]
    second_stage = iv_result["model"]
    add_table_title(doc, "表4  工具变量（FE-2SLS）估计结果")

    rows = [
        ("第一阶段", "lnSubsidy_l1", format_coef(first_stage["coef"], first_stage["f_pval"]), f"Partial F = {first_stage['f_stat']:.2f}", str(iv_result["sample_size"])),
        ("第二阶段", "Overpay", format_coef(second_stage.params["lnSubsidy_l1"], second_stage.pvalues["lnSubsidy_l1"]), f"t = {_get_tstat(second_stage, 'lnSubsidy_l1'):.4f}", str(iv_result["sample_size"])),
    ]

    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["阶段", "因变量", "系数", "统计量", "N"]
    for i, header in enumerate(headers):
        set_cell_text(table.cell(0, i), header, bold=True)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(table.cell(row_idx, col_idx), value, alignment=align)

    apply_three_line_style(table, header_rows=1)
    add_table_note(doc, f"注：第一阶段报告工具变量系数；Partial R² 为 {first_stage['partial_r2']:.4f}，说明工具变量具有统计相关性但解释力度有限。标准误为公司层面聚类稳健标准误。")


def build_subsample_table_docx(doc, df, title, group_specs):
    """分组主回归表（模型2）"""
    control_vars = ["Roa", "Lever", "Top1", "Zone"]
    display_vars = ["lnSubsidy_l1", "Roa", "Lever", "Top1", "Zone"]

    models_data = []
    col_headers = []
    sub_headers = []
    for label, mask in group_specs:
        df_sub = df[mask].copy()
        model, nobs = fit_model(df_sub, "Overpay", ["lnSubsidy_l1"] + control_vars)
        models_data.append((model, nobs))
        col_headers.append(label)
        sub_headers.append("模型2")

    fe = ["控制"] * len(models_data)

    _add_regression_table(
        doc, title, col_headers, sub_headers, display_vars, models_data,
        "注：括号中为 t 值，* p<0.1, ** p<0.05, *** p<0.01。标准误为公司层面聚类稳健标准误。",
        extra_rows=[("Firm FE", fe), ("Year FE", fe)],
    )


def build_robustness_table_docx(doc, df):
    """表6 稳健性检验"""
    control_vars = ["Roa", "Lever", "Top1", "Zone"]
    core_vars = ["lnSubsidy_l1"] + control_vars

    checks = []

    # (1)
    m1, n1 = fit_model(df, "lnCEOpay", core_vars)
    checks.append(("(1) 替换因变量", "lnCEOpay", m1, n1, "lnSubsidy_l1"))

    # (2)
    df_r2 = df[(df["Year"] >= 2010) & (df["Year"] <= 2020)].copy()
    m2, n2 = fit_model(df_r2, "Overpay", core_vars)
    checks.append(("(2) 缩小样本期", "Overpay", m2, n2, "lnSubsidy_l1"))

    # (3)
    df_r4 = df[df["Industry"] == 1].copy()
    m4, n4 = fit_model(df_r4, "Overpay", core_vars)
    checks.append(("(3) 仅制造业", "Overpay", m4, n4, "lnSubsidy_l1"))

    # (4)
    df_alt = df.sort_values(["Symbol", "Year"]).copy()
    if "lnSubsidy1p_l1" not in df_alt.columns:
        df_alt["lnSubsidy1p_l1"] = df_alt.groupby("Symbol")["lnSubsidy1p"].shift(1)
    core_alt = ["lnSubsidy1p_l1"] + control_vars
    m5, n5 = fit_model(df_alt, "Overpay", core_alt)
    checks.append(("(4) 替换解释变量", "Overpay", m5, n5, "lnSubsidy1p_l1"))

    # 构建表格
    add_table_title(doc, "表6  稳健性检验结果")
    n_checks = len(checks)
    # 行：表头2 + 系数1 + t值1 + Controls + Industry FE + Year FE + N + R² = 9
    table = doc.add_table(rows=9, cols=1 + n_checks)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头行1
    set_cell_text(table.cell(0, 0), "变量", bold=True)
    for i, (label, _, _, _, _) in enumerate(checks):
        set_cell_text(table.cell(0, i + 1), label, bold=True, font_size=9)

    # 表头行2 (因变量)
    set_cell_text(table.cell(1, 0), "", bold=True)
    for i, (_, dep, _, _, _) in enumerate(checks):
        set_cell_text(table.cell(1, i + 1), dep, bold=True, font_size=9)

    # 系数行
    set_cell_text(table.cell(2, 0), "补助变量", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    for i, (_, _, model, _, var_name) in enumerate(checks):
        if model and var_name in model.params:
            set_cell_text(table.cell(2, i + 1), format_coef(model.params[var_name], model.pvalues[var_name]))

    # t值行
    set_cell_text(table.cell(3, 0), "")
    for i, (_, _, model, _, var_name) in enumerate(checks):
        if model and var_name in model.params:
            set_cell_text(table.cell(3, i + 1), format_tval(_get_tstat(model, var_name)))

    # Controls, FE, N, R²
    row_labels = ["Controls", "Firm FE", "Year FE", "N", "R²"]
    for r, label in enumerate(row_labels):
        set_cell_text(table.cell(4 + r, 0), label, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        for i, (_, _, model, n, _) in enumerate(checks):
            if label == "N":
                set_cell_text(table.cell(4 + r, i + 1), str(n))
            elif label == "R²":
                set_cell_text(table.cell(4 + r, i + 1), f"{model.rsquared:.4f}" if model else "—")
            else:
                set_cell_text(table.cell(4 + r, i + 1), "控制")

    apply_three_line_style(table, header_rows=2)
    add_table_note(doc, "注：括号中为 t 值，* p<0.1, ** p<0.05, *** p<0.01。标准误为公司层面聚类稳健标准误。")


# ============================================================
# 主流程
# ============================================================

def main():
    data_dir = os.path.join(os.getcwd(), "processed_data")
    print("加载数据...")
    df, _ = build_analysis_dataset(data_dir)
    main_results = run_regressions(df)

    print("\n生成 Word 文档...")
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    # 文档标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("回归分析结果")
    run.font.size = Pt(16)
    _set_run_font(run, "Times New Roman", "黑体")
    run.bold = True

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于数据挖掘的上市公司财政补贴与高管超额薪酬研究")
    run.font.size = Pt(12)
    _set_run_font(run, "Times New Roman", "宋体")

    # 逐表生成
    print("  生成表1 描述性统计...")
    build_desc_table_docx(doc, df)

    print("  生成表2 第一阶段期望薪酬模型...")
    build_first_stage_table_docx(doc, df)

    print("  生成表3 主回归...")
    build_main_table_docx(doc, main_results)

    print("  生成表4 FE-2SLS...")
    build_causal_table_docx(doc, main_results)

    print("  生成表5 中介效应...")
    build_mediation_table_docx(doc, main_results)

    soe_df = df[df["IsSOE"] == 1].copy()
    identified_ratio = soe_df["IsCentralSOE"].notna().mean() if len(soe_df) > 0 else np.nan

    print("  生成表6 稳健性检验...")
    build_robustness_table_docx(doc, df)

    print("  生成表7 分产权性质...")
    build_subsample_table_docx(
        doc, df,
        "表7  分产权性质主回归结果（模型2）",
        [("国有", df["IsSOE"] == 1), ("私营", df["IsPrivate"] == 1)],
    )

    print("  生成表8 分行业管制...")
    build_subsample_table_docx(
        doc, df,
        "表8  分行业管制主回归结果（模型2）",
        [("管制行业", df["RegulatedIndustry"] == 1), ("非管制行业", df["RegulatedIndustry"] == 0)],
    )

    print("  生成表9 央企vs地方国企...")
    build_subsample_table_docx(
        doc, df,
        f"表9  央企与地方国企主回归结果（模型2，央地可识别比例={identified_ratio:.2%}）",
        [("央企", df["IsCentralSOE"] == 1), ("地方国企", df["IsCentralSOE"] == 0)],
    )

    # 保存
    output_dir = os.path.join(os.getcwd(), "results")
    os.makedirs(output_dir, exist_ok=True)
    output_file = os.path.join(output_dir, "回归结果表.docx")
    doc.save(output_file)
    print(f"\n✅ Word 文件已保存至: {output_file}")


if __name__ == "__main__":
    main()
