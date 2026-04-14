"""
使用学校模板骨架 + Pandoc 预处理 + docxcompose 拼装，
稳定导出更符合学院模板要求的论文 Word 文档。
"""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from docxcompose.composer import Composer
from lxml import etree, html as lxml_html
from PIL import Image, ImageDraw, ImageFont


ROOT_DIR = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT_DIR / "thesis_final_bundle" / "thesis_final_humanized_v2.md"
PREPROCESSED_MD = ROOT_DIR / "thesis_final_bundle" / "thesis_final_humanized_v2_pandoc.md"
OUTPUT_DOCX = ROOT_DIR / "thesis_final_bundle" / "thesis_final_humanized_v2.docx"
TABLE_IMAGE_DIR = ROOT_DIR / "thesis_final_bundle" / "generated_tables"
REFERENCE_DOC = ROOT_DIR / "09 【模板 2026届】 论文正文（信息科学与技术学院）.docx"
KAITI_MD = ROOT_DIR / "kaiti.md"
KAITI_REPORT_DOCX = ROOT_DIR / "彭晓俞_开题报告 (1).docx"

COVER_TITLE = "本 科 生 毕 业 论 文 正 文"
DEFAULT_COHORT_TEXT = "（ 2026 届）"
COLLEGE_TEXT = "信息科学与技术学院"
HEADER_LEFT_TEXT = "本科生毕业设计（论文）"
HEADER_SPACER = "                                                                    "
CHINESE_ABSTRACT_FIRST_INDENT_PT = 17.9
TEMPLATE_DEBUG_FILES = (
    ROOT_DIR / "thesis_final_bundle" / "_pandoc_one_shot.docx",
    ROOT_DIR / "thesis_final_bundle" / "_tmp_template_test.docx",
    ROOT_DIR / "thesis_final_bundle" / "_toc_test.docx",
)
FIGURE_TITLE_CROP_TOP = {
    "fig1_lasso_path.png": 80,
    "fig2_rf_importance.png": 70,
    "fig3_shap_summary.png": 65,
    "fig4_shap_subsidy.png": 60,
    "fig4a_shap_subsidy_issoe.png": 60,
    "fig4b_shap_subsidy_mgshder.png": 60,
    "fig8_model_comparison.png": 50,
}


def extract_cover_info_from_kaiti(path: Path) -> dict[str, str]:
    if not path.exists():
        return {}

    text = path.read_text(encoding="utf-8")
    info: dict[str, str] = {}

    title_match = re.search(r"\*\*论文题目\*\*\s*\n([^\n]+)", text)
    if title_match:
        info["thesis_title"] = title_match.group(1).strip()

    year_match = re.search(r"（\s*(\d{4})\s*届\s*）", text)
    if year_match:
        info["year"] = year_match.group(1)

    field_patterns = {
        "student_name": [r"(?:学生姓名|姓名)\s*[:：]\s*([^\s|]+)"],
        "student_id": [r"(?:学号|学\s*号)\s*[:：]\s*([A-Za-z0-9]+)"],
        "major": [r"专业\s*[:：]\s*([^\n|]+)"],
        "class_name": [r"班级\s*[:：]\s*([^\n|]+)"],
        "advisor_name": [r"(?:指导教师|指导老师|导师)\s*[:：]\s*([^\s|]+)"],
        "advisor_title": [r"职称\s*[:：]\s*([^\n|]+)"],
    }

    for key, patterns in field_patterns.items():
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                info[key] = match.group(1).strip()
                break

    return info


def extract_cover_info_from_report_docx(path: Path) -> dict[str, str]:
    if not path.exists():
        return {}

    doc = Document(path)
    info: dict[str, str] = {}

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        year_match = re.search(r"（\s*(\d{4})\s*届\s*）", text)
        if year_match:
            info["year"] = year_match.group(1)

    for table in doc.tables:
        cells = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(cells) >= 3 and len(cells[0]) >= 4:
            if cells[0][0] == "学生姓名" and cells[0][2] == "学号":
                if cells[0][1]:
                    info["student_name"] = cells[0][1]
                if cells[0][3]:
                    info["student_id"] = cells[0][3]
                if cells[1][1]:
                    info["major"] = cells[1][1]
                if cells[1][3]:
                    info["class_name"] = cells[1][3]
                if cells[2][1]:
                    info["advisor_name"] = cells[2][1]
                if cells[2][3]:
                    info["advisor_title"] = cells[2][3]
            if len(cells[0]) >= 2 and cells[0][0] == "论文题目" and cells[0][1]:
                info["thesis_title"] = cells[0][1]
    return {key: value for key, value in info.items() if value}


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", "", text)


def extract_display_equation_tags(text: str) -> list[int | None]:
    tags: list[int | None] = []
    lines = text.splitlines()
    in_block = False
    block_lines: list[str] = []

    for raw_line in lines:
        stripped = raw_line.strip()
        if not in_block:
            if not stripped.startswith("$$"):
                continue
            if stripped.endswith("$$") and stripped != "$$":
                tags.append(extract_equation_tag(stripped[2:-2]))
            elif stripped == "$$":
                in_block = True
                block_lines = []
        else:
            if stripped == "$$":
                tags.append(extract_equation_tag("\n".join(block_lines)))
                in_block = False
                block_lines = []
            else:
                block_lines.append(raw_line)

    return tags


def extract_equation_tag(block_text: str) -> int | None:
    match = re.search(r"\\tag\{(\d+)\}", block_text)
    return int(match.group(1)) if match else None


def collect_markdown_blocks(lines: list[str]) -> list[str]:
    blocks: list[str] = []
    current_block: list[str] = []
    for line in lines:
        if line.strip():
            current_block.append(line.strip())
        elif current_block:
            blocks.append(" ".join(current_block).strip())
            current_block = []
    if current_block:
        blocks.append(" ".join(current_block).strip())
    return blocks


def strip_bold_wrapper(text: str) -> str:
    stripped = text.strip()
    match = re.fullmatch(r"\*\*(.*?)\*\*", stripped)
    return match.group(1).strip() if match else stripped


def extract_title_abstract_keywords_body(text: str) -> tuple[str, list[str], str, str, list[str], str, str]:
    lines = text.splitlines()
    title = next((line[2:].strip() for line in lines if line.startswith("# ")), "")
    if not title:
        raise ValueError("未找到论文标题（# 一级标题）。")

    abstract_heading_idx = next(
        (i for i, line in enumerate(lines) if re.fullmatch(r"\*\*摘\s*要\*\*", line.strip())),
        None,
    )
    keywords_idx = next(
        (i for i, line in enumerate(lines) if re.match(r"^\*\*关键词\*\*[:：]", line.strip())),
        None,
    )
    body_idx = next((i for i, line in enumerate(lines) if re.match(r"^##\s+", line)), None)
    if abstract_heading_idx is None or keywords_idx is None or body_idx is None:
        raise ValueError("未能从 Markdown 中完整解析出摘要、关键词或正文起点。")

    english_heading_idx = next(
        (
            i
            for i in range(keywords_idx + 1, body_idx)
            if re.fullmatch(r"\*\*ABSTRACT\*\*", lines[i].strip(), flags=re.IGNORECASE)
        ),
        None,
    )
    english_keywords_idx = next(
        (
            i
            for i in range((english_heading_idx or keywords_idx) + 1, body_idx)
            if re.match(r"^\*\*Keywords\*\*[:：]", lines[i].strip(), flags=re.IGNORECASE)
        ),
        None,
    )

    abstract_blocks = collect_markdown_blocks(lines[abstract_heading_idx + 1 : keywords_idx])

    keywords_line = lines[keywords_idx].strip()
    keywords = re.sub(r"^\*\*关键词\*\*[:：]\s*", "", keywords_line).strip()

    english_title = ""
    english_abstract_blocks: list[str] = []
    english_keywords = ""
    if english_heading_idx is not None:
        english_title_candidates = [
            strip_bold_wrapper(line)
            for line in lines[keywords_idx + 1 : english_heading_idx]
            if line.strip() and line.strip() != "---"
        ]
        english_title = english_title_candidates[-1] if english_title_candidates else ""
        english_abstract_end_idx = english_keywords_idx or body_idx
        english_abstract_blocks = collect_markdown_blocks(lines[english_heading_idx + 1 : english_abstract_end_idx])
    if english_keywords_idx is not None:
        english_keywords_line = lines[english_keywords_idx].strip()
        english_keywords = re.sub(r"^\*\*Keywords\*\*[:：]\s*", "", english_keywords_line, flags=re.IGNORECASE).strip()

    body = "\n".join(lines[body_idx:]).strip() + "\n"
    return title, abstract_blocks, keywords, english_title, english_abstract_blocks, english_keywords, body


def strip_rules(text: str) -> str:
    lines = [line for line in text.splitlines() if line.strip() != "---"]
    return "\n".join(lines)


def promote_headings(text: str) -> str:
    promoted: list[str] = []
    for line in text.splitlines():
        if re.match(r"^####\s+", line):
            promoted.append(re.sub(r"^####\s+", "### ", line))
        elif re.match(r"^###\s+", line):
            promoted.append(re.sub(r"^###\s+", "## ", line))
        elif re.match(r"^##\s+", line):
            promoted.append(re.sub(r"^##\s+", "# ", line))
        else:
            promoted.append(line)
    return "\n".join(promoted)


def normalize_caption_lines(text: str) -> str:
    normalized: list[str] = []
    for line in text.splitlines():
        match = re.match(r"^\*\*((?:表|图)[^*]+)\*\*$", line.strip())
        if match:
            normalized.append(match.group(1))
        else:
            normalized.append(line)
    return "\n".join(normalized)


def reorder_figure_caption_blocks(text: str) -> str:
    lines = text.splitlines()
    output: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if re.match(r"^图\d", line.strip()):
            j = i + 1
            blank_lines: list[str] = []
            while j < len(lines) and not lines[j].strip():
                blank_lines.append(lines[j])
                j += 1
            if j < len(lines) and re.match(r"^!\[[^\]]*\]\([^)]+\)", lines[j].strip()):
                output.append(lines[j])
                output.extend(blank_lines or [""])
                output.append(line)
                output.append("")
                i = j + 1
                continue
        output.append(line)
        i += 1
    return "\n".join(output)


def _escape_table_cell(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", text).strip()
    cleaned = cleaned.replace("|", r"\|")
    return cleaned if cleaned else " "


def html_table_to_pipe(table_html: str) -> str:
    root = lxml_html.fragment_fromstring(table_html, create_parent="div")
    table = root.xpath(".//table")[0]

    header_rows = table.xpath("./thead/tr")
    body_rows = table.xpath("./tbody/tr") or table.xpath("./tr")
    if not header_rows and body_rows:
        header_rows = [body_rows[0]]
        body_rows = body_rows[1:]

    header_cells = [
        _escape_table_cell(" ".join(cell.itertext()))
        for cell in header_rows[0].xpath("./th|./td")
    ]
    if not header_cells:
        raise ValueError("遇到无法解析的空表格。")

    lines = [
        "| " + " | ".join(header_cells) + " |",
        "| " + " | ".join(["---"] * len(header_cells)) + " |",
    ]

    for row in body_rows:
        cells = [
            _escape_table_cell(" ".join(cell.itertext()))
            for cell in row.xpath("./th|./td")
        ]
        if not cells:
            continue
        if len(cells) < len(header_cells):
            cells.extend([" "] * (len(header_cells) - len(cells)))
        elif len(cells) > len(header_cells):
            cells = cells[: len(header_cells)]
        lines.append("| " + " | ".join(cells) + " |")

    return "\n".join(lines)


def table_image_widths(column_count: int, header_text: str) -> list[float]:
    if column_count == 8:
        return [2.4, 1.8, 1.7, 1.5, 1.2, 1.2, 1.5, 1.0]
    if column_count == 7:
        return [4.2, 1.6, 1.7, 1.7, 1.7, 1.7, 1.7]
    if column_count == 6:
        return [3.6, 2.6, 2.6, 1.6, 1.8, 1.4]
    if column_count == 5:
        return [3.2, 3.7, 3.3, 2.9, 1.8]
    if column_count == 4:
        return [3.8, 3.4, 3.4, 3.4]
    if column_count == 3:
        if "定义" in header_text:
            return [3.7, 2.7, 8.6]
        if "排名" in header_text:
            return [2.0, 6.0, 4.0]
        return [4.0, 5.0, 5.0]
    if column_count == 2:
        return [6.0, 6.0]
    return [1.0] * max(column_count, 1)


def load_table_font(size: int, *, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/STHeiti Medium.ttc" if bold else "/System/Library/Fonts/Songti.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def text_width(draw: ImageDraw.ImageDraw, text: str, font) -> int:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0]


def wrap_table_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return [""]
    lines: list[str] = []
    current = ""
    for char in text:
        candidate = current + char
        if current and text_width(draw, candidate, font) > max_width:
            lines.append(current)
            current = char
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines or [""]


def parse_html_table_rows(table_html: str) -> tuple[list[list[dict[str, object]]], int]:
    root = lxml_html.fragment_fromstring(table_html, create_parent="div")
    table = root.xpath(".//table")[0]
    parsed_rows: list[list[dict[str, object]]] = []
    column_count = 0
    for row in table.xpath(".//tr"):
        cells = []
        row_span = 0
        for cell in row.xpath("./th|./td"):
            text = re.sub(r"\s+", " ", " ".join(cell.itertext())).strip()
            colspan_match = re.search(r"\d+", cell.get("colspan", "1") or "1")
            colspan = int(colspan_match.group(0)) if colspan_match else 1
            cells.append({"text": text, "colspan": colspan})
            row_span += colspan
        if cells:
            parsed_rows.append(cells)
            column_count = max(column_count, row_span)
    return parsed_rows, column_count


def render_html_table_image(table_html: str, output_path: Path) -> None:
    rows, column_count = parse_html_table_rows(table_html)
    if not rows or column_count == 0:
        raise ValueError("遇到无法解析的空表格。")

    canvas_width = 1800
    margin_x = 42
    margin_y = 28
    padding_x = 12
    padding_y = 9
    font = load_table_font(34)
    bold_font = load_table_font(34, bold=True)
    line_height = 48
    drawable = Image.new("RGB", (canvas_width, 200), "white")
    draw = ImageDraw.Draw(drawable)

    ratios = table_image_widths(column_count, " ".join(str(cell["text"]) for cell in rows[0]))
    total_ratio = sum(ratios)
    table_width = canvas_width - margin_x * 2
    column_widths = [int(table_width * ratio / total_ratio) for ratio in ratios]
    column_widths[-1] += table_width - sum(column_widths)

    wrapped_rows: list[list[dict[str, object]]] = []
    row_heights: list[int] = []
    for r_idx, row in enumerate(rows):
        wrapped_row = []
        cursor = 0
        row_height = 0
        for cell in row:
            colspan = int(cell["colspan"])
            span_width = sum(column_widths[cursor : cursor + colspan])
            cell_font = bold_font if r_idx == 0 else font
            lines = wrap_table_text(draw, str(cell["text"]), cell_font, span_width - padding_x * 2)
            wrapped_row.append({"lines": lines, "colspan": colspan})
            row_height = max(row_height, len(lines) * line_height + padding_y * 2)
            cursor += colspan
        wrapped_rows.append(wrapped_row)
        row_heights.append(max(row_height, line_height + padding_y * 2))

    canvas_height = margin_y * 2 + sum(row_heights) + 8
    image = Image.new("RGB", (canvas_width, canvas_height), "white")
    draw = ImageDraw.Draw(image)
    x0 = margin_x
    y = margin_y
    rule_width = 3

    draw.line((x0, y, x0 + table_width, y), fill="black", width=rule_width)
    for r_idx, row in enumerate(wrapped_rows):
        cursor = 0
        row_top = y
        row_height = row_heights[r_idx]
        for cell in row:
            colspan = int(cell["colspan"])
            span_width = sum(column_widths[cursor : cursor + colspan])
            lines = cell["lines"]
            cell_font = bold_font if r_idx == 0 else font
            block_height = len(lines) * line_height
            text_y = row_top + (row_height - block_height) / 2
            left = x0 + sum(column_widths[:cursor])
            is_first_column = cursor == 0
            for line in lines:
                line_width = text_width(draw, line, cell_font)
                if is_first_column and colspan == 1:
                    text_x = left + padding_x
                else:
                    text_x = left + (span_width - line_width) / 2
                draw.text((text_x, text_y), line, fill="black", font=cell_font)
                text_y += line_height
            cursor += colspan
        y += row_height
        if r_idx == 0:
            draw.line((x0, y, x0 + table_width, y), fill="black", width=rule_width)
    draw.line((x0, y, x0 + table_width, y), fill="black", width=rule_width)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path)


def convert_html_tables(text: str) -> str:
    pattern = re.compile(r"<table\b[^>]*>.*?</table>", re.DOTALL | re.IGNORECASE)
    table_index = 0

    def replace_table(match: re.Match[str]) -> str:
        nonlocal table_index
        table_index += 1
        image_path = TABLE_IMAGE_DIR / f"table_{table_index:02d}.png"
        render_html_table_image(match.group(0), image_path)
        rel_path = image_path.relative_to(SOURCE_MD.parent).as_posix()
        return f"\n\n![]({rel_path}){{width=15.5cm}}\n\n"

    return pattern.sub(replace_table, text)


def superscript_body_citations_markdown(text: str) -> str:
    lines = text.splitlines()
    converted: list[str] = []
    in_references = False
    for line in lines:
        if re.match(r"^#\s+参考文献\s*$", line.strip()):
            in_references = True
            converted.append(line)
            continue
        if in_references:
            converted.append(re.sub(r"^(\[\d+\])\s+", r"\1", line))
            continue
        converted.append(re.sub(r"(?<!\$)\[(\d{1,2})\]", r"<sup>[\1]</sup>", line))
    return "\n".join(converted)


def build_pandoc_body_markdown(body_text: str) -> str:
    if TABLE_IMAGE_DIR.exists():
        for old_image in TABLE_IMAGE_DIR.glob("table_*.png"):
            old_image.unlink()
    TABLE_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    body = strip_rules(body_text)
    body = normalize_caption_lines(body)
    body = reorder_figure_caption_blocks(body)
    body = promote_headings(body)
    body = superscript_body_citations_markdown(body)
    body = convert_html_tables(body)
    body = re.sub(r"\n{3,}", "\n\n", body).strip()
    return body + "\n"


def build_pandoc_abstract_markdown(
    abstract_blocks: list[str],
    keywords: str,
    english_title: str,
    english_abstract_blocks: list[str],
    english_keywords: str,
) -> str:
    parts = [block for block in abstract_blocks if block]
    if keywords:
        parts.append(f"关键词：{keywords}")
    if english_title:
        parts.append(english_title)
    if english_abstract_blocks:
        parts.append("ABSTRACT")
        parts.extend(block for block in english_abstract_blocks if block)
    if english_keywords:
        parts.append(f"Keywords: {english_keywords}")
    return "\n\n".join(parts).strip() + "\n"


def _set_run_font(run, western: str = "Times New Roman", east_asia: str = "宋体") -> None:
    run.font.name = western
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), western)
    r_fonts.set(qn("w:hAnsi"), western)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_text(paragraph, text: str) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    if text:
        paragraph.add_run(text)


def apply_font(
    paragraph,
    size_pt: float,
    *,
    bold: bool | None = None,
    east_asia: str = "宋体",
    western: str = "Times New Roman",
) -> None:
    for run in paragraph.runs:
        if bold is not None:
            run.bold = bold
        run.font.size = Pt(size_pt)
        _set_run_font(run, western=western, east_asia=east_asia)


def apply_style_font(
    style,
    size_pt: float,
    *,
    bold: bool | None = None,
    east_asia: str = "宋体",
    western: str = "Times New Roman",
) -> None:
    if bold is not None:
        style.font.bold = bold
    style.font.size = Pt(size_pt)
    style.font.name = western
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), western)
    r_fonts.set(qn("w:hAnsi"), western)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_cover_cell_text(cell, text: str, *, size_pt: float = 12, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.3
    set_paragraph_text(paragraph, text)
    apply_font(paragraph, size_pt, east_asia="楷体", western="楷体")


def set_cover_title_cell_text(cell, text: str, *, size_pt: float = 14) -> None:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        lines = [""]

    for paragraph in cell.paragraphs[1:]:
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)

    first_paragraph = cell.paragraphs[0]
    target_paragraphs = [first_paragraph]
    while len(target_paragraphs) < len(lines):
        target_paragraphs.append(cell.add_paragraph())

    for idx, paragraph in enumerate(target_paragraphs):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(lines) > 1 else WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        set_paragraph_text(paragraph, lines[idx] if idx < len(lines) else "")
        apply_font(paragraph, size_pt, east_asia="楷体", western="楷体")


def fill_cover_cell_text(cell, text: str, *, size_pt: float = 12) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(paragraph, text)
    apply_font(paragraph, size_pt, east_asia="楷体", western="楷体")


def split_cover_title(title: str, max_len: int = 18) -> tuple[str, str]:
    compact = re.sub(r"\s+", "", title)
    if len(compact) <= max_len:
        return compact, ""
    for token in ("与", "和", "及", "：", ":", "——"):
        idx = compact.find(token)
        if 8 <= idx <= max_len + 2:
            return compact[:idx], compact[idx:]
    mid = len(compact) // 2
    return compact[:mid], compact[mid:]


def format_chinese_title_lines(title: str) -> str:
    line1, line2 = split_cover_title(title)
    return f"{line1}\n{line2}" if line2 else line1


def split_english_title(title: str, max_len: int = 42) -> list[str]:
    words = title.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if current and len(candidate) > max_len:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines


def element_tag(element) -> str:
    return element.tag.split("}")[-1]


def element_text(element) -> str:
    return "".join(element.itertext()).strip()


def has_section_break(element) -> bool:
    if element_tag(element) == "sectPr":
        return True
    return bool(element.xpath('.//*[local-name()="sectPr"]'))


def is_empty_paragraph_element(element) -> bool:
    return element_tag(element) == "p" and not element_text(element).strip()


def is_abstract_heading(text: str) -> bool:
    normalized = normalize_text(text)
    return normalized != "" and set(normalized) <= {"摘", "要"} and "摘" in normalized and "要" in normalized


def is_toc_heading(text: str) -> bool:
    normalized = normalize_text(text)
    return normalized != "" and set(normalized) <= {"目", "录"} and "目" in normalized and "录" in normalized


def remove_template_shapes(doc: Document) -> None:
    # Keep template artwork such as the school logo, but remove orange instruction text boxes.
    removable = []
    for textbox in doc.element.xpath('.//*[local-name()="txbxContent"]'):
        node = textbox
        while node is not None and element_tag(node) not in {"drawing", "pict"}:
            node = node.getparent()
        removable.append(node if node is not None else textbox)

    for node in removable:
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)


def locate_template_layout(doc: Document) -> dict[str, int]:
    blocks = list(doc.element.body.iterchildren())
    section_breaks = [
        idx
        for idx, element in enumerate(blocks)
        if element_tag(element) == "p" and has_section_break(element)
    ]
    if len(section_breaks) < 3:
        raise ValueError("学校模板的分节结构异常，无法定位封面、摘要、目录和正文边界。")

    abstract_heading_idx = next(
        idx
        for idx in range(section_breaks[0] + 1, section_breaks[1])
        if element_tag(blocks[idx]) == "p" and is_abstract_heading(element_text(blocks[idx]))
    )
    toc_title_idx = next(
        idx
        for idx in range(section_breaks[1] + 1, section_breaks[2])
        if element_tag(blocks[idx]) == "p" and is_toc_heading(element_text(blocks[idx]))
    )

    abstract_title_idx = next(
        idx
        for idx in range(abstract_heading_idx - 1, section_breaks[0], -1)
        if element_tag(blocks[idx]) == "p" and normalize_text(element_text(blocks[idx])) != ""
    )

    abstract_insert_idx = abstract_heading_idx + 1
    if abstract_insert_idx < section_breaks[1] and is_empty_paragraph_element(blocks[abstract_insert_idx]):
        abstract_insert_idx += 1

    toc_insert_idx = toc_title_idx + 1
    if toc_insert_idx < section_breaks[2] and is_empty_paragraph_element(blocks[toc_insert_idx]):
        toc_insert_idx += 1

    return {
        "cover_section_break_idx": section_breaks[0],
        "abstract_section_break_idx": section_breaks[1],
        "toc_section_break_idx": section_breaks[2],
        "abstract_title_idx": abstract_title_idx,
        "abstract_heading_idx": abstract_heading_idx,
        "abstract_insert_idx": abstract_insert_idx,
        "toc_title_idx": toc_title_idx,
        "toc_insert_idx": toc_insert_idx,
        "body_insert_idx": section_breaks[2] + 1,
    }


def remove_block_range(doc: Document, start_idx: int, end_idx: int) -> None:
    blocks = list(doc.element.body.iterchildren())
    for element in blocks[start_idx:end_idx]:
        doc.element.body.remove(element)


def remove_paragraph(paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def set_paragraph_basic_format(paragraph, *, line_spacing: float = 1.3, first_indent: float = 0) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(first_indent)


def style_toc_styles(doc: Document) -> None:
    indent_by_level = {1: 0, 2: 21, 3: 42}
    for style in doc.styles:
        style_id = (getattr(style, "style_id", "") or "").lower().replace(" ", "")
        style_name = (style.name or "").lower().replace(" ", "")
        level = None
        for candidate in (1, 2, 3):
            if style_id == f"toc{candidate}" or style_name in {
                f"toc{candidate}",
                f"目录{candidate}",
            }:
                level = candidate
                break
        if level is None:
            continue
        style.paragraph_format.line_spacing = 1.3
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.left_indent = Pt(indent_by_level[level])
        style.paragraph_format.first_line_indent = Pt(0)
        apply_style_font(style, 10.5, bold=False, east_asia="宋体", western="Times New Roman")


def is_toc_paragraph(paragraph) -> bool:
    style_id = (getattr(paragraph.style, "style_id", "") or "").lower().replace(" ", "")
    style_name = (getattr(paragraph.style, "name", "") or "").lower().replace(" ", "")
    return style_id in {"toc1", "toc2", "toc3"} or style_name in {
        "toc1",
        "toc2",
        "toc3",
        "toc 1",
        "toc 2",
        "toc 3",
        "目录1",
        "目录2",
        "目录3",
    }


def extract_existing_toc_page_cache(path: Path) -> dict[str, str]:
    if not path.exists():
        return {}

    doc = Document(path)
    page_cache: dict[str, str] = {}
    for paragraph in doc.paragraphs:
        if not is_toc_paragraph(paragraph):
            continue
        text = paragraph.text.strip()
        if "\t" not in text:
            continue
        title, page = text.rsplit("\t", 1)
        title = re.sub(r"\s+", " ", title.strip())
        page = page.strip()
        if title and page:
            page_cache[title] = page
    return page_cache


def render_docx_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    profile_dir = output_dir / "_lo_profile"
    subprocess.run(
        [
            "soffice",
            f"-env:UserInstallation=file://{profile_dir}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists():
        raise FileNotFoundError(f"未生成 PDF：{pdf_path}")
    return pdf_path


def build_pdf_toc_page_cache(docx_path: Path, titles: list[str]) -> dict[str, str]:
    if not titles or shutil.which("soffice") is None or shutil.which("pdfinfo") is None or shutil.which("pdftotext") is None:
        return {}

    with tempfile.TemporaryDirectory(prefix="toc_page_cache_") as tmp_dir_name:
        tmp_dir = Path(tmp_dir_name)
        pdf_path = render_docx_to_pdf(docx_path, tmp_dir)
        pdfinfo = subprocess.run(
            ["pdfinfo", str(pdf_path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        ).stdout
        match = re.search(r"Pages:\s+(\d+)", pdfinfo)
        if not match:
            return {}

        total_pages = int(match.group(1))
        normalized_titles = {normalize_text(title): title for title in titles if title.strip()}
        page_cache: dict[str, str] = {}

        for pdf_page in range(1, total_pages + 1):
            page_text = subprocess.run(
                ["pdftotext", "-f", str(pdf_page), "-l", str(pdf_page), str(pdf_path), "-"],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            ).stdout
            footer_match = re.search(r"第\s*(\d+)\s*页", page_text)
            if not footer_match:
                continue

            logical_page = footer_match.group(1)
            normalized_page_text = normalize_text(page_text)
            for normalized_title, raw_title in normalized_titles.items():
                if raw_title in page_cache or not normalized_title:
                    continue
                if normalized_title in normalized_page_text:
                    page_cache[raw_title] = logical_page

        return page_cache


def toc_level_for_heading(text: str) -> int | None:
    if re.match(r"^第[一二三四五六七八九十]+章\s+", text) or text in {"参考文献", "致谢"}:
        return 1
    if is_heading_2(text):
        return 2
    if is_heading_3(text):
        return None
    return None


def max_bookmark_id(doc: Document) -> int:
    bookmark_ids = []
    for value in doc.element.xpath('.//*[local-name()="bookmarkStart"]/@*[local-name()="id"]'):
        try:
            bookmark_ids.append(int(value))
        except ValueError:
            continue
    return max(bookmark_ids, default=0)


def remove_auto_toc_bookmarks(doc: Document) -> None:
    auto_ids: set[str] = set()
    for start in doc.element.xpath('.//*[local-name()="bookmarkStart"]'):
        name = start.get(qn("w:name"))
        if name and name.startswith("_TocAuto"):
            bookmark_id = start.get(qn("w:id"))
            if bookmark_id is not None:
                auto_ids.add(bookmark_id)
            parent = start.getparent()
            if parent is not None:
                parent.remove(start)

    for end in doc.element.xpath('.//*[local-name()="bookmarkEnd"]'):
        bookmark_id = end.get(qn("w:id"))
        if bookmark_id in auto_ids:
            parent = end.getparent()
            if parent is not None:
                parent.remove(end)


def add_heading_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    insert_idx = 0
    if paragraph._p.pPr is not None:
        insert_idx = paragraph._p.index(paragraph._p.pPr) + 1
    paragraph._p.insert(insert_idx, start)
    paragraph._p.append(end)


def add_body_end_bookmark(doc: Document, name: str = "_BodyEnd") -> None:
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip():
            add_heading_bookmark(paragraph, name, max_bookmark_id(doc) + 1)
            return
    raise ValueError("未找到可用于标记正文末页的段落。")


def insert_section_break_before(paragraph, sect_pr_template) -> None:
    new_paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    sect_pr = deepcopy(sect_pr_template)
    for child in list(sect_pr):
        if element_tag(child) == "pgNumType":
            sect_pr.remove(child)
    p_pr.append(sect_pr)
    new_paragraph.append(p_pr)
    paragraph._p.addprevious(new_paragraph)


def split_body_for_back_matter(doc: Document) -> None:
    references_para = next((p for p in doc.paragraphs if p.text.strip() == "参考文献"), None)
    acknowledgements_para = next((p for p in doc.paragraphs if p.text.strip() == "致谢"), None)
    if references_para is None or acknowledgements_para is None:
        raise ValueError("未找到参考文献或致谢标题，无法拆分页眉分节。")

    references_para.paragraph_format.page_break_before = False
    acknowledgements_para.paragraph_format.page_break_before = False

    body_sect_template = deepcopy(doc.sections[-1]._sectPr)
    insert_section_break_before(references_para, body_sect_template)
    insert_section_break_before(acknowledgements_para, body_sect_template)


def set_section_header_text(section, right_text: str) -> None:
    section.header.is_linked_to_previous = False
    paragraphs = section.header.paragraphs
    header_paragraph = paragraphs[0] if paragraphs else section.header.add_paragraph()
    for paragraph in paragraphs[1:]:
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)
    header_paragraph.style = section.part.document.styles["Header"]
    set_paragraph_text(header_paragraph, f"{HEADER_LEFT_TEXT}{HEADER_SPACER}{right_text} ")
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def configure_back_matter_headers(doc: Document) -> None:
    if len(doc.sections) < 6:
        raise ValueError(f"分节数量不足，当前仅 {len(doc.sections)} 节，无法单独设置参考文献和致谢页眉。")

    set_section_header_text(doc.sections[3], "论文正文")
    set_section_header_text(doc.sections[4], "参考文献")
    set_section_header_text(doc.sections[5], "致谢")


def _append_toc_run(parent, text: str = "", *, hyperlink_style: bool = False, web_hidden: bool = False):
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    if hyperlink_style:
        r_style = OxmlElement("w:rStyle")
        r_style.set(qn("w:val"), "af")
        r_pr.append(r_style)
    no_proof = OxmlElement("w:noProof")
    r_pr.append(no_proof)
    if web_hidden:
        r_pr.append(OxmlElement("w:webHidden"))
    run.append(r_pr)
    if text:
        node = OxmlElement("w:t")
        if text != text.strip():
            node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        node.text = text
        run.append(node)
    parent.append(run)
    return run


def _append_toc_field(
    parent,
    instr: str,
    result_text: str | None = None,
    *,
    result_web_hidden: bool = False,
) -> None:
    begin_run = _append_toc_run(parent, web_hidden=True)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)

    instr_run = _append_toc_run(parent, web_hidden=True)
    instr_node = OxmlElement("w:instrText")
    instr_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_node.text = f" {instr} "
    instr_run.append(instr_node)

    sep_run = _append_toc_run(parent, web_hidden=True)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run.append(sep)

    if result_text is not None:
        _append_toc_run(parent, result_text, web_hidden=result_web_hidden)

    end_run = _append_toc_run(parent, web_hidden=True)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)


def _append_toc_field_begin(parent, instr: str) -> None:
    begin_run = _append_toc_run(parent, web_hidden=True)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)

    instr_run = _append_toc_run(parent, web_hidden=True)
    instr_node = OxmlElement("w:instrText")
    instr_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_node.text = f" {instr} "
    instr_run.append(instr_node)

    sep_run = _append_toc_run(parent, web_hidden=True)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run.append(sep)


def _append_toc_field_end(parent) -> None:
    end_run = _append_toc_run(parent, web_hidden=True)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)


def build_template_toc_paragraph(
    title: str,
    level: int,
    bookmark_name: str,
    page_result: str,
    *,
    start_toc_field: bool = False,
    end_toc_field: bool = False,
):
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    p_style = OxmlElement("w:pStyle")
    p_style.set(qn("w:val"), f"TOC{level}")
    p_pr.append(p_style)
    paragraph.append(p_pr)

    if start_toc_field:
        _append_toc_field_begin(paragraph, r'TOC \o "1-3" \h \z \u')

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")
    for part in re.split(r"(\s+)", title):
        if part:
            _append_toc_run(hyperlink, part, hyperlink_style=True)

    tab_run = _append_toc_run(hyperlink, web_hidden=True)
    tab_run.append(OxmlElement("w:tab"))
    _append_toc_field(
        hyperlink,
        f"PAGEREF {bookmark_name} \\h",
        page_result,
        result_web_hidden=True,
    )
    paragraph.append(hyperlink)

    if end_toc_field:
        _append_toc_field_end(paragraph)
    return paragraph


def heading_entries_for_template_toc(doc: Document) -> list[dict[str, str | int]]:
    entries: list[dict[str, str | int]] = []
    next_id = max_bookmark_id(doc) + 1
    remove_auto_toc_bookmarks(doc)

    layout = locate_template_layout(doc)
    blocks = list(doc.element.body.iterchildren())
    para_map = {paragraph._p: paragraph for paragraph in doc.paragraphs}
    for block in blocks[layout["body_insert_idx"] :]:
        paragraph = para_map.get(block)
        if paragraph is None:
            continue
        text = re.sub(r"\s+", " ", paragraph.text.strip())
        level = toc_level_for_heading(text)
        if level is None:
            continue

        bookmark_name = f"_TocAuto{len(entries) + 1:04d}"
        add_heading_bookmark(paragraph, bookmark_name, next_id)
        next_id += 1
        entries.append({"title": text, "level": level, "bookmark": bookmark_name})

    return entries


def replace_toc_with_template_fields(doc: Document, page_cache: dict[str, str]) -> None:
    entries = heading_entries_for_template_toc(doc)
    if not entries:
        raise ValueError("未找到正文标题，无法按模板生成目录。")

    layout = locate_template_layout(doc)
    remove_block_range(doc, layout["toc_insert_idx"], layout["toc_section_break_idx"])

    layout = locate_template_layout(doc)
    for offset, entry in enumerate(entries):
        title = str(entry["title"])
        page_result = page_cache.get(title, "1")
        paragraph = build_template_toc_paragraph(
            title=title,
            level=int(entry["level"]),
            bookmark_name=str(entry["bookmark"]),
            page_result=page_result,
            start_toc_field=(offset == 0),
            end_toc_field=(offset == len(entries) - 1),
        )
        doc.element.body.insert(layout["toc_insert_idx"] + offset, paragraph)


def style_cover_and_titles(doc: Document, thesis_title: str, cover_info: dict[str, str], cohort_text: str) -> None:
    paragraphs = doc.paragraphs

    cover_title_para = next((p for p in paragraphs[:10] if "毕业论文正文" in normalize_text(p.text)), None)
    if cover_title_para is not None:
        set_paragraph_text(cover_title_para, COVER_TITLE)
        cover_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_font(cover_title_para, 26, bold=True, east_asia="黑体", western="黑体")

    cohort_para = next((p for p in paragraphs[:12] if "届" in p.text), None)
    if cohort_para is not None:
        set_paragraph_text(cohort_para, cohort_text)
        cohort_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_font(cohort_para, 16, east_asia="楷体", western="楷体")

    college_para = next((p for p in paragraphs[:20] if "信息科学与技术学院" in p.text), None)
    if college_para is not None:
        set_paragraph_text(college_para, COLLEGE_TEXT)
        college_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_font(college_para, 16, bold=True, east_asia="仿宋", western="仿宋")

    cover_table = doc.tables[0]
    cover_title = re.sub(r"\s+", "", cover_info.get("thesis_title", thesis_title))
    cover_title_text = format_chinese_title_lines(cover_title)
    set_cover_title_cell_text(
        cover_table.cell(0, 1),
        cover_title_text,
        size_pt=14,
    )
    fill_cover_cell_text(cover_table.cell(1, 1), "", size_pt=12)
    fill_cover_cell_text(cover_table.cell(2, 0), "", size_pt=12)
    fill_cover_cell_text(cover_table.cell(3, 1), cover_info.get("student_name", ""))
    fill_cover_cell_text(cover_table.cell(3, 3), cover_info.get("student_id", ""))
    fill_cover_cell_text(cover_table.cell(4, 1), cover_info.get("major", ""))
    fill_cover_cell_text(cover_table.cell(4, 3), cover_info.get("class_name", ""))
    fill_cover_cell_text(cover_table.cell(5, 1), cover_info.get("advisor_name", ""))
    fill_cover_cell_text(cover_table.cell(5, 3), cover_info.get("advisor_title", ""))


def style_abstract_and_toc_shell(doc: Document, layout: dict[str, int], thesis_title: str) -> None:
    blocks = list(doc.element.body.iterchildren())
    para_map = {paragraph._p: paragraph for paragraph in doc.paragraphs}

    abstract_title_para = para_map[blocks[layout["abstract_title_idx"]]]
    set_paragraph_text(abstract_title_para, format_chinese_title_lines(thesis_title))
    abstract_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_basic_format(abstract_title_para)
    apply_font(abstract_title_para, 16, east_asia="黑体", western="黑体")

    abstract_heading_para = para_map[blocks[layout["abstract_heading_idx"]]]
    set_paragraph_text(abstract_heading_para, "摘  要")
    abstract_heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_basic_format(abstract_heading_para)
    apply_font(abstract_heading_para, 14, east_asia="黑体", western="黑体")

    toc_title_para = para_map[blocks[layout["toc_title_idx"]]]
    set_paragraph_text(toc_title_para, "目  录")
    toc_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_basic_format(toc_title_para)
    apply_font(toc_title_para, 16, east_asia="黑体", western="黑体")


def style_keywords_paragraph(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, "")
    label, _, content = text.partition("：")
    label_run = paragraph.add_run(label + "：")
    content_run = paragraph.add_run(content)
    paragraph.style = paragraph.part.document.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_basic_format(paragraph, first_indent=0)
    paragraph.paragraph_format.left_indent = Pt(48)
    paragraph.paragraph_format.first_line_indent = Pt(-48)
    label_run.font.size = Pt(10.5)
    label_run.bold = None
    _set_run_font(label_run, western="黑体", east_asia="黑体")
    content_run.font.size = Pt(10.5)
    _set_run_font(content_run, western="Times New Roman", east_asia="宋体")


def style_english_keywords_paragraph(paragraph, text: str) -> None:
    set_paragraph_text(paragraph, "")
    label, _, content = text.partition(":")
    label_run = paragraph.add_run(label + ":")
    content_run = paragraph.add_run(content)
    paragraph.style = paragraph.part.document.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_basic_format(paragraph, first_indent=0)
    paragraph.paragraph_format.left_indent = Pt(57)
    paragraph.paragraph_format.first_line_indent = Pt(-57)
    label_run.font.size = Pt(10.5)
    label_run.bold = True
    _set_run_font(label_run, western="Times New Roman", east_asia="Times New Roman")
    content_run.font.size = Pt(10.5)
    _set_run_font(content_run, western="Times New Roman", east_asia="Times New Roman")


def insert_paragraph_after(
    paragraph,
    *,
    text: str = "",
    line_spacing: float = 1.0,
    alignment=None,
    first_indent: float = 0,
):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = paragraph.part.document.styles["Normal"]
    new_para.alignment = paragraph.alignment if alignment is None else alignment
    fmt = new_para.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(first_indent)
    if text:
        new_para.add_run(text)
    return new_para


def insert_blank_paragraph_after(paragraph, *, line_spacing: float = 1.0):
    return insert_paragraph_after(paragraph, line_spacing=line_spacing)


def normalize_blank_paragraphs(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            continue
        if has_section_break(paragraph._p):
            continue

        paragraph.style = doc.styles["Normal"]
        fmt = paragraph.paragraph_format
        fmt.line_spacing = 1.0
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.first_line_indent = Pt(0)

        if not paragraph.runs:
            paragraph.add_run("")

        for run in paragraph.runs:
            run.font.size = Pt(10.5)
            _set_run_font(run, western="Times New Roman", east_asia="宋体")


def set_cell_borders(cell, *, top: bool = False, bottom: bool = False) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for name in ("top", "bottom", "left", "right", "insideH", "insideV"):
        existing = borders.find(qn(f"w:{name}"))
        if existing is not None:
            borders.remove(existing)
    if top:
        el = OxmlElement("w:top")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    if bottom:
        el = OxmlElement("w:bottom")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:color"), "000000")
        borders.append(el)


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    rows = table.rows
    if not rows:
        return
    set_table_column_widths(table)

    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = Cm(_table_column_widths(len(row.cells), row)[c_idx])
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                set_paragraph_basic_format(paragraph, line_spacing=1.0, first_indent=0)
                apply_font(paragraph, 9)
                if r_idx == 0:
                    for run in paragraph.runs:
                        run.bold = True
            if r_idx == 0:
                set_cell_borders(cell, top=True, bottom=True)
            elif r_idx == len(rows) - 1:
                set_cell_borders(cell, bottom=True)


def _table_column_widths(column_count: int, row=None) -> list[float]:
    header_text = ""
    if row is not None:
        header_text = " ".join(cell.text.strip() for cell in row.cells)
    if column_count == 8:
        return [2.5, 2.0, 1.8, 1.6, 1.2, 1.2, 1.6, 1.2]
    if column_count == 7:
        return [4.3, 1.7, 1.8, 1.8, 1.8, 1.8, 1.8]
    if column_count == 6:
        return [3.4, 2.7, 2.4, 1.6, 1.8, 1.4]
    if column_count == 5:
        return [3.0, 3.4, 3.1, 3.0, 1.8]
    if column_count == 4:
        return [3.8, 3.4, 3.4, 3.4]
    if column_count == 3:
        if "定义" in header_text:
            return [4.0, 2.6, 8.6]
        if "排名" in header_text:
            return [2.0, 6.0, 4.0]
        return [4.0, 5.0, 5.0]
    if column_count == 2:
        return [6.0, 6.0]
    return [max(1.2, 15.0 / max(column_count, 1))] * column_count


def set_table_column_widths(table) -> None:
    column_count = len(table.columns)
    if column_count == 0:
        return
    widths_cm = _table_column_widths(column_count, table.rows[0])
    widths_twips = [str(int(width / 2.54 * 1440)) for width in widths_cm]

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(width) for width in widths_twips)))

    grid = tbl.tblGrid
    if grid is not None:
        for grid_col in list(grid):
            grid.remove(grid_col)
        for width in widths_twips:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), width)
            grid.append(grid_col)

    for row in table.rows:
        row_cells = row.cells
        for c_idx, cell in enumerate(row_cells):
            if c_idx >= len(widths_twips):
                continue
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), widths_twips[c_idx])


def postprocess_abstract_docx(path: Path) -> None:
    doc = Document(path)
    in_english_abstract = False
    english_title_first_para = None
    english_title_last_para = None
    abstract_heading_para = None
    chinese_body_para = None
    chinese_keywords_para = None
    english_body_para = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith("关键词：") or text.startswith("关键词:"):
            style_keywords_paragraph(paragraph, text.replace("关键词:", "关键词："))
            chinese_keywords_para = paragraph
            continue
        if text.upper() == "ABSTRACT":
            in_english_abstract = True
            abstract_heading_para = paragraph
            paragraph.style = doc.styles["Normal"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_basic_format(paragraph, first_indent=0)
            apply_font(paragraph, 14, bold=True, east_asia="Times New Roman", western="Times New Roman")
            continue
        if text.startswith("Keywords:") or text.startswith("Keywords："):
            style_english_keywords_paragraph(paragraph, text.replace("Keywords：", "Keywords:"))
            continue
        if re.fullmatch(r"[A-Z0-9 ,:&'\-()]+", text) and len(text) > 20:
            title_lines = split_english_title(text)
            set_paragraph_text(paragraph, title_lines[0])
            paragraph.style = doc.styles["Normal"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_basic_format(paragraph, first_indent=0)
            paragraph.paragraph_format.page_break_before = False
            apply_font(paragraph, 16, bold=True, east_asia="Times New Roman", western="Times New Roman")
            english_title_first_para = paragraph
            prev_para = paragraph
            for line in title_lines[1:]:
                prev_para = insert_paragraph_after(
                    prev_para,
                    text=line,
                    line_spacing=1.3,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    first_indent=0,
                )
                apply_font(prev_para, 16, bold=True, east_asia="Times New Roman", western="Times New Roman")
            english_title_last_para = prev_para
            continue
        paragraph.style = doc.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if in_english_abstract:
            set_paragraph_basic_format(paragraph, first_indent=21)
            apply_font(paragraph, 10.5, east_asia="Times New Roman", western="Times New Roman")
            english_body_para = paragraph
        else:
            set_paragraph_basic_format(paragraph, first_indent=CHINESE_ABSTRACT_FIRST_INDENT_PT)
            apply_font(paragraph, 10.5)
            chinese_body_para = paragraph

    if chinese_body_para is not None:
        insert_blank_paragraph_after(chinese_body_para, line_spacing=1.0)
    if chinese_keywords_para is not None:
        insert_blank_paragraph_after(chinese_keywords_para, line_spacing=1.0)
    if chinese_keywords_para is not None and english_title_first_para is not None:
        english_page_blank = insert_paragraph_after(
            chinese_keywords_para,
            line_spacing=1.0,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_indent=0,
        )
        english_page_blank.paragraph_format.page_break_before = True
        apply_font(english_page_blank, 10.5, east_asia="宋体", western="Times New Roman")
        english_page_blank_2 = insert_paragraph_after(
            english_page_blank,
            line_spacing=1.0,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_indent=0,
        )
        apply_font(english_page_blank_2, 10.5, east_asia="宋体", western="Times New Roman")
    if english_title_last_para is not None:
        insert_blank_paragraph_after(english_title_last_para, line_spacing=1.0)
    if abstract_heading_para is not None:
        insert_blank_paragraph_after(abstract_heading_para, line_spacing=1.0)
    if english_body_para is not None:
        insert_blank_paragraph_after(english_body_para, line_spacing=1.0)
    normalize_blank_paragraphs(doc)
    doc.save(path)


def is_heading_2(text: str) -> bool:
    return bool(re.match(r"^\d+\.\d+\s+", text))


def is_heading_3(text: str) -> bool:
    return bool(re.match(r"^\d+\.\d+\.\d+\s+", text))


def is_caption_text(text: str) -> bool:
    return bool(re.match(r"^(表|图)\d+(?:-\d+)?[A-Za-z]?\s+[^。！？]+$", text))


def _build_text_run(text: str, r_pr_template=None, *, superscript: bool = False):
    run = OxmlElement("w:r")
    if r_pr_template is not None:
        run.append(deepcopy(r_pr_template))
    if superscript:
        r_pr = run.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            run.insert(0, r_pr)
        vert_align = r_pr.find(qn("w:vertAlign"))
        if vert_align is None:
            vert_align = OxmlElement("w:vertAlign")
            r_pr.append(vert_align)
        vert_align.set(qn("w:val"), "superscript")

    text_node = OxmlElement("w:t")
    if text != text.strip():
        text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_node.text = text
    run.append(text_node)
    return run


def superscript_citations(paragraph) -> None:
    citation_pattern = re.compile(r"\[(\d{1,2})\]")
    for run in list(paragraph.runs):
        text = run.text
        if not citation_pattern.search(text):
            continue

        parent = run._r.getparent()
        if parent is None:
            continue
        insert_idx = parent.index(run._r)
        r_pr_template = run._r.find(qn("w:rPr"))
        new_runs = []
        last_idx = 0
        for match in citation_pattern.finditer(text):
            if match.start() > last_idx:
                new_runs.append(_build_text_run(text[last_idx : match.start()], r_pr_template))
            new_runs.append(_build_text_run(match.group(0), r_pr_template, superscript=True))
            last_idx = match.end()
        if last_idx < len(text):
            new_runs.append(_build_text_run(text[last_idx:], r_pr_template))

        for offset, new_run in enumerate(new_runs):
            parent.insert(insert_idx + offset, new_run)
        parent.remove(run._r)


def add_equation_number(paragraph, number: int) -> None:
    paragraph.style = paragraph.part.document.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_basic_format(paragraph, first_indent=0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Cm(15.2),
        alignment=WD_TAB_ALIGNMENT.RIGHT,
        leader=WD_TAB_LEADER.SPACES,
    )

    tab_run = paragraph.add_run()
    tab_run.add_tab()
    tab_run.font.size = Pt(10.5)
    _set_run_font(tab_run, western="Times New Roman", east_asia="Times New Roman")

    left_paren = paragraph.add_run("(")
    left_paren.font.size = Pt(10.5)
    _set_run_font(left_paren, western="宋体", east_asia="宋体")

    num_run = paragraph.add_run(str(number))
    num_run.font.size = Pt(10.5)
    _set_run_font(num_run, western="Times New Roman", east_asia="Times New Roman")

    right_paren = paragraph.add_run(")")
    right_paren.font.size = Pt(10.5)
    _set_run_font(right_paren, western="宋体", east_asia="宋体")


def postprocess_body_docx(path: Path, equation_tags: list[int | None] | None = None) -> None:
    doc = Document(path)
    heading3_style = doc.styles["标题3"] if "标题3" in doc.styles else doc.styles["Heading 3"]
    first_chapter_seen = False
    in_chapter6 = False
    in_references = False
    display_equation_idx = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "oMathPara" in paragraph._p.xml:
            if equation_tags and display_equation_idx < len(equation_tags):
                equation_number = equation_tags[display_equation_idx]
                if equation_number is not None:
                    add_equation_number(paragraph, equation_number)
            display_equation_idx += 1
            continue

        if not text:
            continue

        if re.match(r"^第[一二三四五六七八九十]+章\s+", text) or text in {"参考文献", "致谢"}:
            paragraph.style = doc.styles["Heading 1"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_basic_format(paragraph)
            paragraph.paragraph_format.page_break_before = first_chapter_seen
            apply_font(paragraph, 16, bold=True, east_asia="黑体")
            first_chapter_seen = True
            in_chapter6 = text.startswith("第六章")
            in_references = text == "参考文献"
            continue

        if is_heading_2(text):
            paragraph.style = doc.styles["Heading 2"]
            paragraph.style.font.bold = False
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_basic_format(paragraph)
            apply_font(paragraph, 14, bold=False, east_asia="黑体")
            for run in paragraph.runs:
                run.bold = False
            continue

        if is_heading_3(text):
            paragraph.style = heading3_style
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_basic_format(paragraph, first_indent=21)
            apply_font(paragraph, 12, east_asia="宋体")
            for run in paragraph.runs:
                run.bold = False
            continue

        if is_caption_text(text):
            paragraph.style = doc.styles["Normal"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_basic_format(paragraph)
            apply_font(paragraph, 9, bold=True)
            continue

        if text.startswith("注：") or text.startswith("注:"):
            paragraph.style = doc.styles["Normal"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_basic_format(paragraph, first_indent=0)
            apply_font(paragraph, 10.5, bold=False)
            continue

        if in_references:
            if re.match(r"^\[\d+\]\s+", text):
                set_paragraph_text(paragraph, re.sub(r"^(\[\d+\])\s+", r"\1", text))
            paragraph.style = doc.styles["Normal"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_basic_format(paragraph, first_indent=0)
            apply_font(paragraph, 10.5, bold=False)
            continue

        paragraph.style = doc.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_basic_format(paragraph, first_indent=21)
        preserve_bold = in_chapter6 and bool(re.match(r"^(?:\d+\.\s+|（\d+）)", text))
        apply_font(paragraph, 10.5, bold=None if preserve_bold else False)
        if not in_references:
            superscript_citations(paragraph)

    for table in doc.tables:
        style_table(table)

    normalize_blank_paragraphs(doc)
    doc.save(path)


def run_pandoc(
    input_md: Path,
    output_docx: Path,
    *,
    extra_resource_paths: list[str] | None = None,
) -> None:
    pandoc = shutil.which("pandoc")
    if pandoc is None:
        raise RuntimeError("未找到 pandoc，请先安装 pandoc。")

    resource_paths = [
        str(SOURCE_MD.parent),
        str(ROOT_DIR / "results"),
        str(ROOT_DIR),
    ]
    if extra_resource_paths:
        resource_paths = extra_resource_paths + resource_paths
    cmd = [
        pandoc,
        "--from",
        "markdown+tex_math_dollars+raw_tex+raw_html+pipe_tables",
        str(input_md),
        "-o",
        str(output_docx),
        "--reference-doc",
        str(REFERENCE_DOC),
        "--resource-path=" + ":".join(resource_paths),
        "--wrap=none",
    ]
    completed = subprocess.run(cmd, capture_output=True, text=True)
    if completed.returncode != 0:
        raise RuntimeError(completed.stderr.strip() or "Pandoc 导出失败。")


def prepare_docx_export_images(asset_root: Path) -> None:
    try:
        from PIL import Image
    except ImportError as exc:  # pragma: no cover - environment guard
        raise RuntimeError("缺少 Pillow，无法在导出时处理论文图片。") from exc

    image_dir = asset_root / "images"
    image_dir.mkdir(parents=True, exist_ok=True)
    source_dir = SOURCE_MD.parent / "images"

    for source_path in source_dir.glob("*"):
        target_path = image_dir / source_path.name
        if source_path.name in FIGURE_TITLE_CROP_TOP and source_path.suffix.lower() in {".png", ".jpg", ".jpeg"}:
            crop_top = FIGURE_TITLE_CROP_TOP[source_path.name]
            with Image.open(source_path) as image:
                cropped = image.crop((0, crop_top, image.size[0], image.size[1]))
                cropped.save(target_path)
        else:
            shutil.copy2(source_path, target_path)


def insert_doc_body_elements(target_doc: Document, insert_idx: int, source_doc: Document) -> None:
    elements = [deepcopy(element) for element in source_doc.element.body if element_tag(element) != "sectPr"]
    for offset, element in enumerate(elements):
        target_doc.element.body.insert(insert_idx + offset, element)


def enable_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def rewrite_update_fields_setting(docx_path: Path) -> None:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        with zipfile.ZipFile(docx_path) as zf:
            zf.extractall(tmp_path)

        settings_xml = tmp_path / "word" / "settings.xml"
        settings_root = etree.parse(str(settings_xml)).getroot()
        update_fields = settings_root.find("w:updateFields", namespaces=ns)
        if update_fields is None:
            update_fields = etree.SubElement(settings_root, f"{{{ns['w']}}}updateFields")
        update_fields.set(f"{{{ns['w']}}}val", "true")
        etree.ElementTree(settings_root).write(
            str(settings_xml),
            encoding="UTF-8",
            xml_declaration=True,
            standalone="yes",
        )

        rebuilt_docx = tmp_path / "rebuilt.docx"
        with zipfile.ZipFile(rebuilt_docx, "w", zipfile.ZIP_DEFLATED) as zf:
            for file_path in sorted(tmp_path.rglob("*")):
                if file_path == rebuilt_docx or file_path.is_dir():
                    continue
                zf.write(file_path, file_path.relative_to(tmp_path))
        shutil.move(rebuilt_docx, docx_path)


def count_referenced_images(docx_path: Path) -> int:
    with zipfile.ZipFile(docx_path) as zf:
        count = 0
        for name in zf.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            try:
                root = etree.fromstring(zf.read(name))
            except etree.XMLSyntaxError:
                continue
            count += len(root.xpath('.//*[local-name()="blip"]/@*[local-name()="embed"]'))
            count += len(root.xpath('.//*[local-name()="imagedata"]/@*[local-name()="id"]'))
    return count


def rewrite_body_footer_page_fields(docx_path: Path) -> None:
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    w_ns = ns["w"]
    r_ns = ns["r"]

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        with zipfile.ZipFile(docx_path) as zf:
            zf.extractall(tmp_path)

        document_xml = tmp_path / "word" / "document.xml"
        rels_xml = tmp_path / "word" / "_rels" / "document.xml.rels"

        document_root = etree.parse(str(document_xml))
        rels_root = etree.parse(str(rels_xml))

        rel_map = {
            rel.get("Id"): rel.get("Target")
            for rel in rels_root.xpath("/*[local-name()='Relationships']/*[local-name()='Relationship']")
        }

        sect_prs = document_root.xpath("//*[local-name()='sectPr']")
        if not sect_prs:
            raise ValueError("未找到文档分节信息，无法修正正文页码。")

        body_sect = sect_prs[-1]
        footer_refs = body_sect.xpath("./w:footerReference[@w:type='default']", namespaces=ns)
        if not footer_refs:
            raise ValueError("正文分节未找到默认页脚引用，无法修正正文页码。")

        footer_rid = footer_refs[0].get(f"{{{r_ns}}}id")
        footer_target = rel_map.get(footer_rid)
        if not footer_target:
            raise ValueError("无法解析正文页脚关系目标。")

        footer_xml = tmp_path / "word" / footer_target
        footer_root = etree.parse(str(footer_xml)).getroot()

        for child in list(footer_root):
            footer_root.remove(child)

        p = etree.SubElement(footer_root, f"{{{w_ns}}}p")
        p_pr = etree.SubElement(p, f"{{{w_ns}}}pPr")
        p_style = etree.SubElement(p_pr, f"{{{w_ns}}}pStyle")
        p_style.set(f"{{{w_ns}}}val", "a7")
        jc = etree.SubElement(p_pr, f"{{{w_ns}}}jc")
        jc.set(f"{{{w_ns}}}val", "center")
        p_rpr = etree.SubElement(p_pr, f"{{{w_ns}}}rPr")
        p_lang = etree.SubElement(p_rpr, f"{{{w_ns}}}lang")
        p_lang.set(f"{{{w_ns}}}val", "zh-CN")

        def append_text_run(text: str, *, east_asia: bool = True) -> None:
            r = etree.SubElement(p, f"{{{w_ns}}}r")
            r_pr = etree.SubElement(r, f"{{{w_ns}}}rPr")
            if east_asia:
                r_fonts = etree.SubElement(r_pr, f"{{{w_ns}}}rFonts")
                r_fonts.set(f"{{{w_ns}}}hint", "eastAsia")
            lang = etree.SubElement(r_pr, f"{{{w_ns}}}lang")
            lang.set(f"{{{w_ns}}}val", "zh-CN")
            t = etree.SubElement(r, f"{{{w_ns}}}t")
            if text != text.strip():
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = text

        def append_field(instr: str, result_text: str) -> None:
            r_begin = etree.SubElement(p, f"{{{w_ns}}}r")
            r_begin_pr = etree.SubElement(r_begin, f"{{{w_ns}}}rPr")
            begin_lang = etree.SubElement(r_begin_pr, f"{{{w_ns}}}lang")
            begin_lang.set(f"{{{w_ns}}}val", "zh-CN")
            fld_begin = etree.SubElement(r_begin, f"{{{w_ns}}}fldChar")
            fld_begin.set(f"{{{w_ns}}}fldCharType", "begin")
            fld_begin.set(f"{{{w_ns}}}dirty", "true")

            r_instr = etree.SubElement(p, f"{{{w_ns}}}r")
            r_instr_pr = etree.SubElement(r_instr, f"{{{w_ns}}}rPr")
            instr_lang = etree.SubElement(r_instr_pr, f"{{{w_ns}}}lang")
            instr_lang.set(f"{{{w_ns}}}val", "zh-CN")
            instr_text = etree.SubElement(r_instr, f"{{{w_ns}}}instrText")
            instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            instr_text.text = f" {instr} "

            r_sep = etree.SubElement(p, f"{{{w_ns}}}r")
            r_sep_pr = etree.SubElement(r_sep, f"{{{w_ns}}}rPr")
            sep_lang = etree.SubElement(r_sep_pr, f"{{{w_ns}}}lang")
            sep_lang.set(f"{{{w_ns}}}val", "zh-CN")
            fld_sep = etree.SubElement(r_sep, f"{{{w_ns}}}fldChar")
            fld_sep.set(f"{{{w_ns}}}fldCharType", "separate")

            r_result = etree.SubElement(p, f"{{{w_ns}}}r")
            r_result_pr = etree.SubElement(r_result, f"{{{w_ns}}}rPr")
            result_lang = etree.SubElement(r_result_pr, f"{{{w_ns}}}lang")
            result_lang.set(f"{{{w_ns}}}val", "zh-CN")
            result_t = etree.SubElement(r_result, f"{{{w_ns}}}t")
            result_t.text = result_text

            r_end = etree.SubElement(p, f"{{{w_ns}}}r")
            r_end_pr = etree.SubElement(r_end, f"{{{w_ns}}}rPr")
            end_lang = etree.SubElement(r_end_pr, f"{{{w_ns}}}lang")
            end_lang.set(f"{{{w_ns}}}val", "zh-CN")
            fld_end = etree.SubElement(r_end, f"{{{w_ns}}}fldChar")
            fld_end.set(f"{{{w_ns}}}fldCharType", "end")

        append_text_run("第")
        append_field("PAGE", "1")
        append_text_run("页")
        append_text_run("  ", east_asia=False)
        append_text_run("共")
        append_field("SECTIONPAGES", "1")
        append_text_run("页")

        etree.ElementTree(footer_root).write(str(footer_xml), encoding="utf-8", xml_declaration=True)

        with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as out_zip:
            for path in sorted(tmp_path.rglob("*")):
                if path.is_file():
                    out_zip.write(path, path.relative_to(tmp_path))


def build_scaffold_doc(
    thesis_title: str,
    cohort_text: str,
    cover_info: dict[str, str],
    abstract_docx: Path,
) -> Document:
    scaffold = Document(REFERENCE_DOC)
    remove_template_shapes(scaffold)
    layout = locate_template_layout(scaffold)
    style_abstract_and_toc_shell(scaffold, layout, thesis_title)

    remove_block_range(
        scaffold,
        layout["abstract_insert_idx"],
        layout["abstract_section_break_idx"],
    )

    layout = locate_template_layout(scaffold)
    remove_block_range(
        scaffold,
        layout["toc_insert_idx"],
        layout["toc_section_break_idx"],
    )

    layout = locate_template_layout(scaffold)
    final_body_idx = len(list(scaffold.element.body.iterchildren())) - 1
    remove_block_range(scaffold, layout["body_insert_idx"], final_body_idx)

    layout = locate_template_layout(scaffold)
    insert_doc_body_elements(scaffold, layout["abstract_insert_idx"], Document(abstract_docx))
    enable_update_fields(scaffold)
    style_toc_styles(scaffold)
    return scaffold


def cleanup_old_docx() -> list[str]:
    removed: list[str] = []
    for path in TEMPLATE_DEBUG_FILES:
        if path.exists():
            path.unlink()
            removed.append(str(path))
    return removed


def main() -> int:
    if not REFERENCE_DOC.exists():
        print(f"未找到学校模板：{REFERENCE_DOC}", file=sys.stderr)
        return 1

    cover_info = extract_cover_info_from_kaiti(KAITI_MD)
    cover_info.update(extract_cover_info_from_report_docx(KAITI_REPORT_DOCX))

    source_text = SOURCE_MD.read_text(encoding="utf-8")
    equation_tags = extract_display_equation_tags(source_text)
    (
        title_from_md,
        abstract_blocks,
        keywords,
        english_title,
        english_abstract_blocks,
        english_keywords,
        body_text,
    ) = extract_title_abstract_keywords_body(source_text)
    cover_info["thesis_title"] = title_from_md
    thesis_title = title_from_md
    cohort_text = f"（ {cover_info['year']} 届）" if cover_info.get("year") else DEFAULT_COHORT_TEXT

    body_md = build_pandoc_body_markdown(body_text)
    PREPROCESSED_MD.write_text(body_md, encoding="utf-8")
    abstract_md = build_pandoc_abstract_markdown(
        abstract_blocks,
        keywords,
        english_title,
        english_abstract_blocks,
        english_keywords,
    )

    tmp_dir = ROOT_DIR / "tmp" / "docs" / "export_school_docx"
    if tmp_dir.exists():
        shutil.rmtree(tmp_dir)
    tmp_dir.mkdir(parents=True, exist_ok=True)

    abstract_md_path = tmp_dir / "abstract.md"
    abstract_docx_path = tmp_dir / "abstract.docx"
    body_docx_path = tmp_dir / "body.docx"
    scaffold_docx_path = tmp_dir / "scaffold.docx"
    asset_root = tmp_dir / "assets"

    abstract_md_path.write_text(abstract_md, encoding="utf-8")
    prepare_docx_export_images(asset_root)

    try:
        extra_resource_paths = [str(asset_root)]
        toc_page_cache = extract_existing_toc_page_cache(OUTPUT_DOCX)

        run_pandoc(abstract_md_path, abstract_docx_path, extra_resource_paths=extra_resource_paths)
        postprocess_abstract_docx(abstract_docx_path)

        run_pandoc(PREPROCESSED_MD, body_docx_path, extra_resource_paths=extra_resource_paths)
        postprocess_body_docx(body_docx_path, equation_tags=equation_tags)

        scaffold = build_scaffold_doc(
            thesis_title=thesis_title,
            cohort_text=cohort_text,
            cover_info=cover_info,
            abstract_docx=abstract_docx_path,
        )
        scaffold.save(scaffold_docx_path)

        composer = Composer(Document(scaffold_docx_path))
        composer.append(Document(body_docx_path))
        composer.save(str(OUTPUT_DOCX))

        final_doc = Document(OUTPUT_DOCX)
        enable_update_fields(final_doc)
        replace_toc_with_template_fields(final_doc, toc_page_cache)
        style_toc_styles(final_doc)
        final_doc.save(OUTPUT_DOCX)

        add_body_end_bookmark(final_doc)
        final_doc.save(OUTPUT_DOCX)

        rewrite_body_footer_page_fields(OUTPUT_DOCX)
        rewrite_update_fields_setting(OUTPUT_DOCX)

        removed = cleanup_old_docx()
        shutil.rmtree(tmp_dir)

        doc = Document(OUTPUT_DOCX)
        image_count = count_referenced_images(OUTPUT_DOCX)
        print(f"导出完成：{OUTPUT_DOCX}")
        print(f"Pandoc 预处理稿：{PREPROCESSED_MD}")
        print(f"Section 数量：{len(doc.sections)}")
        print(f"Word 表格数量：{len(doc.tables)}")
        print(f"文档可见图片数量：{image_count}")
        if removed:
            print("已删除旧测试文档：")
            for path in removed:
                print(f"  - {path}")
        if cover_info:
            print("已填入的封面字段：")
            for key in (
                "thesis_title",
                "year",
                "student_name",
                "student_id",
                "major",
                "class_name",
                "advisor_name",
                "advisor_title",
            ):
                if cover_info.get(key):
                    print(f"  - {key}: {cover_info[key]}")
        return 0
    except Exception as exc:  # pragma: no cover - CLI guard
        print(str(exc), file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
