"""
将 results/数据分析部分_整理.md 转为 results/数据分析部分_整理.docx
使用 python-docx，支持标题层级、段落、Markdown 表格、粗体、行内代码
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement


# ─── 字体辅助 ────────────────────────────────────────────────────────────────

def _set_run_font(run, western="Times New Roman", east="宋体", size_pt=None, bold=None):
    run.font.name = western
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), east)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _add_run(para, text, western="Times New Roman", east="宋体",
             size_pt=10.5, bold=False, italic=False, color=None):
    run = para.add_run(text)
    _set_run_font(run, western, east, size_pt, bold)
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


# ─── 三线表工具 ────────────────────────────────────────────────────────────────

def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val["val"])
        el.set(qn("w:sz"), val["sz"])
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), val.get("color", "000000"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _apply_three_line(table, n_header_rows=1):
    """顶粗线 + 表头下细线 + 底粗线"""
    thick = {"val": "single", "sz": "12"}
    thin  = {"val": "single", "sz": "6"}
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    for c in range(n_cols):
        _set_cell_border(table.cell(0, c), top=thick)
    for c in range(n_cols):
        _set_cell_border(table.cell(n_header_rows - 1, c), bottom=thin)
    for c in range(n_cols):
        _set_cell_border(table.cell(n_rows - 1, c), bottom=thick)


def _set_cell_text(cell, text, bold=False, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after  = Pt(1)
    pf.line_spacing = Pt(14)
    _add_run(p, text, size_pt=size_pt, bold=bold)


# ─── 内联 Markdown 解析（粗体 **...**） ────────────────────────────────────────

def _render_inline(para, raw_text, size_pt=10.5):
    """解析行内 **bold**，其余为普通文本"""
    parts = re.split(r"(\*\*[^*]+\*\*)", raw_text)
    for part in parts:
        m = re.fullmatch(r"\*\*([^*]+)\*\*", part)
        if m:
            _add_run(para, m.group(1), size_pt=size_pt, bold=True)
        else:
            if part:
                _add_run(para, part, size_pt=size_pt)


# ─── Markdown 表格解析 ────────────────────────────────────────────────────────

def _parse_md_table(lines):
    """
    返回 (header_row, data_rows, col_count)
    lines: 以 | 开头的连续行列表
    """
    rows = []
    for line in lines:
        if re.match(r"^\s*\|[-:\s|]+\|\s*$", line):
            continue  # 分隔行
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return [], [], 0
    col_count = max(len(r) for r in rows)
    # 补齐
    rows = [r + [""] * (col_count - len(r)) for r in rows]
    return rows[0], rows[1:], col_count


def _insert_md_table(doc, header, data_rows, col_count):
    """插入三线表"""
    n_rows = 1 + len(data_rows)
    table = doc.add_table(rows=n_rows, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 列宽平均分配（可按需调整）
    # 第1列稍宽
    total_cm = 14.0
    first_col_cm = 4.0
    rest_cm = (total_cm - first_col_cm) / max(col_count - 1, 1) if col_count > 1 else total_cm
    for i, col in enumerate(table.columns):
        width = Cm(first_col_cm if i == 0 else rest_cm)
        for cell in col.cells:
            cell.width = width

    # 表头
    for c, h in enumerate(header):
        raw = re.sub(r"[\*`]", "", h).strip()
        align = WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_text(table.cell(0, c), raw, bold=True, size_pt=9.5, align=align)

    # 数据
    for r, row in enumerate(data_rows):
        for c, val in enumerate(row):
            raw = re.sub(r"[\*`\\]", "", val).strip()
            align = WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_text(table.cell(r + 1, c), raw, size_pt=9.5, align=align)

    _apply_three_line(table, n_header_rows=1)

    # 表格后间距
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)


# ─── 文档样式工具 ────────────────────────────────────────────────────────────

def _set_doc_defaults(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "宋体")

    # 页边距 2.5cm 四周
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)


def _add_heading(doc, text, level):
    """level 1/2/3 → 大/中/小标题"""
    p = doc.add_paragraph()
    sizes   = {1: 16, 2: 14, 3: 12}
    east    = {1: "黑体", 2: "黑体", 3: "黑体"}
    spacing = {1: (18, 12), 2: (14, 8), 3: (10, 6)}
    sz = sizes.get(level, 11)
    ef = east.get(level, "黑体")
    bef, aft = spacing.get(level, (8, 4))
    pf = p.paragraph_format
    pf.space_before = Pt(bef)
    pf.space_after  = Pt(aft)
    _add_run(p, text, western="Times New Roman", east=ef, size_pt=sz, bold=True)


def _add_normal_para(doc, text, indent_cm=0):
    """添加正文段落，支持行内粗体"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(4)
    pf.line_spacing = Pt(18)
    if indent_cm:
        pf.first_line_indent = Cm(indent_cm)
    _render_inline(p, text, size_pt=10.5)


def _add_blockquote(doc, text):
    """引用块 > ..."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent  = Cm(1)
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    pf.line_spacing = Pt(16)
    run = _add_run(p, text.lstrip("> ").strip(), size_pt=9.5)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


# ─── 主转换逻辑 ──────────────────────────────────────────────────────────────

def convert_md_to_docx(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    _set_doc_defaults(doc)

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip("\n")

        # ── 跳过空行 ──────────────────────────────────────────
        if raw.strip() == "":
            i += 1
            continue

        # ── 水平分割线 ---
        if re.match(r"^---+\s*$", raw.strip()):
            i += 1
            continue

        # ── 标题 # ────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)", raw)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            _add_heading(doc, text, level)
            i += 1
            continue

        # ── 引用块 > ──────────────────────────────────────────
        if raw.startswith(">"):
            # 收集连续引用行
            quote_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                quote_lines.append(lines[i].rstrip("\n").lstrip("> ").strip())
                i += 1
            _add_blockquote(doc, " ".join(quote_lines))
            continue

        # ── Markdown 表格 ─────────────────────────────────────
        if raw.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            header, data, ncols = _parse_md_table(table_lines)
            if ncols > 0:
                _insert_md_table(doc, header, data, ncols)
            continue

        # ── 列表项 - / 数字. ────────────────────────────────
        m_li = re.match(r"^(\s*)(?:-|\d+\.)\s+(.*)", raw)
        if m_li:
            indent = len(m_li.group(1))
            text = m_li.group(2).strip()
            p = doc.add_paragraph(style="List Bullet")
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after  = Pt(2)
            pf.left_indent  = Cm(0.5 + indent * 0.3)
            _render_inline(p, text, size_pt=10.5)
            i += 1
            continue

        # ── 普通段落 ──────────────────────────────────────────
        # 收集到空行或结构行为止
        para_lines = []
        while i < len(lines):
            l = lines[i].rstrip("\n")
            if (l.strip() == "" or l.startswith("#") or l.startswith("|")
                    or l.startswith(">") or re.match(r"^---+\s*$", l.strip())
                    or re.match(r"^(\s*)(?:-|\d+\.)\s+", l)):
                break
            para_lines.append(l.strip())
            i += 1
        text = " ".join(para_lines)
        if text:
            _add_normal_para(doc, text, indent_cm=0.74)  # 首行缩进2字符

    doc.save(docx_path)
    print(f"✅ 已保存: {docx_path}")


if __name__ == "__main__":
    base = os.getcwd()
    md_path   = os.path.join(base, "results", "数据分析部分_整理.md")
    docx_path = os.path.join(base, "results", "数据分析部分_整理.docx")
    convert_md_to_docx(md_path, docx_path)
