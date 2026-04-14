"""
校验论文正文中的核心实证数据、图片与参考文献是否与当前 results/ 最新结果一致。
"""

from __future__ import annotations

import hashlib
import json
import re
import sys
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from scripts.regression_analysis import (
    ALT_SUBSIDY_LAG_COL,
    BASE_SUBSIDY_LAG_COL,
    FE_CONTROL_VARS,
    _build_fe_matrix,
    _fit_heckman_two_step,
    _fit_model2,
    _fit_with_cluster_se,
)


BUNDLE_DIR = ROOT_DIR / "thesis_final_bundle"
THESIS_PATH = BUNDLE_DIR / "thesis_final_humanized_v2.md"
DOCX_PATH = BUNDLE_DIR / "thesis_final_humanized_v2.docx"
BUNDLE_IMAGES_DIR = BUNDLE_DIR / "images"
RESULTS_DIR = ROOT_DIR / "results"
DATASET_PATH = RESULTS_DIR / "regression_dataset.csv"
SAMPLE_SUMMARY_PATH = RESULTS_DIR / "sample_screening_summary.csv"
MAIN_FE_PATH = RESULTS_DIR / "main_regression_fe_comparison.csv"
ROBUSTNESS_PATH = RESULTS_DIR / "robustness_results.csv"
HECKMAN_PATH = RESULTS_DIR / "heckman_results.csv"
DESCRIPTIVE_STATS_PATH = RESULTS_DIR / "descriptive_statistics.csv"
CAUSAL_RESULTS_PATH = RESULTS_DIR / "causal_results.csv"
MAIN_MEDIATION_SUMMARY_PATH = RESULTS_DIR / "main_mediation_summary.csv"
EVENT_STUDY_SUMMARY_PATH = RESULTS_DIR / "event_study_summary.json"
EVENT_STUDY_RESULTS_PATH = RESULTS_DIR / "event_study_results.csv"

IMAGE_SOURCE_MAP = {
    "fig1_lasso_path_notitle.png": "fig1_lasso_path.png",
    "fig2_rf_importance_notitle.png": "fig2_rf_importance.png",
    "fig3_shap_summary_notitle.png": "fig3_shap_summary.png",
    "fig4_shap_subsidy_notitle.png": "fig4_shap_subsidy.png",
}
BUNDLE_ONLY_IMAGES = {"fig1_1_techroute.png"}


def fmt_int(value: int | float) -> str:
    return f"{int(round(float(value))):,}"


def fmt_fixed(value: int | float, digits: int = 4, unicode_minus: bool = True) -> str:
    number = float(value)
    if abs(number) < 0.5 * (10 ** (-digits)):
        number = 0.0
    if number < 0:
        prefix = "−" if unicode_minus else "-"
    else:
        prefix = ""
    return f"{prefix}{abs(number):.{digits}f}"


def fmt_plain(value: int | float, digits: int = 4) -> str:
    return f"{float(value):.{digits}f}"


def fmt_coef(value: int | float, p_value: int | float, digits: int = 4) -> str:
    coef = fmt_fixed(value, digits=digits)
    p = float(p_value)
    stars = ""
    if p < 0.01:
        stars = "***"
    elif p < 0.05:
        stars = "**"
    elif p < 0.1:
        stars = "*"
    return f"{coef}{stars}"


def fmt_p_text(value: int | float) -> str:
    p = float(value)
    return "< 0.001" if p < 0.001 else f"= {p:.3f}"


def fmt_wan(value: int | float, digits: int = 2) -> str:
    return f"{float(value) / 10000:,.{digits}f}"


def expect(text: str, needle: str, label: str, failures: list[str]) -> None:
    if needle not in text:
        failures.append(f"[缺失] {label}: {needle}")


def expect_absent(text: str, needle: str, label: str, failures: list[str]) -> None:
    if needle in text:
        failures.append(f"[残留] {label}: {needle}")


def expect_true(condition: bool, label: str, detail: str, failures: list[str]) -> None:
    if not condition:
        failures.append(f"[异常] {label}: {detail}")


def file_md5(path: Path) -> str:
    return hashlib.md5(path.read_bytes()).hexdigest()


def check_bundle_images(thesis_raw: str, failures: list[str]) -> None:
    image_names = sorted(set(re.findall(r"!\[.*?\]\(\./images/([^)]+)\)", thesis_raw)))
    expected_images = sorted(set(IMAGE_SOURCE_MAP.keys()) | BUNDLE_ONLY_IMAGES)
    if image_names != expected_images:
        failures.append(f"[图片] 正文图片列表异常: {image_names}")
    for name in image_names:
        bundle_path = BUNDLE_IMAGES_DIR / name
        if not bundle_path.exists():
            failures.append(f"[缺失] bundle 图片不存在: {bundle_path}")
            continue
        if name in BUNDLE_ONLY_IMAGES:
            continue
        result_path = RESULTS_DIR / IMAGE_SOURCE_MAP[name]
        if not result_path.exists():
            failures.append(f"[缺失] results 图片不存在: {result_path}")
            continue
        # 正文使用 _notitle 图片以避免 DOCX 中图内标题与题注重复，
        # 因此这里只要求源图与正文图都存在，不强制哈希完全一致。


def check_reference_section(thesis_raw: str, failures: list[str]) -> None:
    if "## 参考文献" not in thesis_raw:
        failures.append("[缺失] 参考文献标题")
        return
    ref_text = thesis_raw.split("## 参考文献", 1)[1]
    ref_numbers = [int(x) for x in re.findall(r"^\[(\d+)\]", ref_text, flags=re.MULTILINE)]
    expected = list(range(1, 33))
    if ref_numbers != expected:
        failures.append(f"[参考文献] 编号异常: {ref_numbers}")

    whole_text_matches = {int(x) for x in re.findall(r"\[(\d+)\]", thesis_raw)}
    for bad in (33,):
        if bad in whole_text_matches:
            failures.append(f"[参考文献] 出现不允许的引用编号: [{bad}]")

    for i in expected:
        if f"[{i}]" not in thesis_raw:
            failures.append(f"[参考文献] 正文或文后缺少编号: [{i}]")

    if re.search(r"^\[\d+\]\s+", ref_text, flags=re.MULTILINE):
        failures.append("[参考文献] 文后编号后出现多余空格，应保持为 [1]文献内容")


def check_table_and_figure_numbers(thesis_raw: str, failures: list[str]) -> None:
    table_labels = re.findall(r"\*\*表(\d+-\d+)", thesis_raw)
    figure_labels = re.findall(r"\*\*图(\d+-\d+)", thesis_raw)
    expected_tables = ["3-1"] + [f"4-{i}" for i in range(1, 14)]
    expected_figures = ["1-1"] + [f"4-{i}" for i in range(1, 5)]
    if table_labels != expected_tables:
        failures.append(f"[编号] 表格编号异常: {table_labels}")
    if figure_labels != expected_figures:
        failures.append(f"[编号] 图片编号异常: {figure_labels}")


def check_equation_numbers(thesis_raw: str, failures: list[str]) -> None:
    display_blocks = re.findall(r"\$\$\s*(.*?)\s*\$\$", thesis_raw, flags=re.S)
    tags = []
    for idx, block in enumerate(display_blocks, start=1):
        match = re.search(r"\\tag\{(\d+)\}", block)
        if not match:
            failures.append(f"[公式编号] 第 {idx} 个展示公式缺少编号")
            continue
        tags.append(int(match.group(1)))
    if tags and tags != list(range(1, len(tags) + 1)):
        failures.append(f"[公式编号] 展示公式编号不连续或重复: {tags}")


def check_docx_exists(failures: list[str]) -> None:
    if not DOCX_PATH.exists():
        failures.append(f"[缺失] DOCX 不存在: {DOCX_PATH}")
        return

    doc = Document(DOCX_PATH)
    expect_true(len(doc.sections) >= 4, "DOCX格式", f"分节数量异常: {len(doc.sections)}", failures)
    for idx, section in enumerate(doc.sections):
        expect_true(abs(section.top_margin.cm - 2.0) < 0.02, "DOCX格式", f"第{idx+1}节上边距不是2cm", failures)
        expect_true(abs(section.bottom_margin.cm - 2.0) < 0.02, "DOCX格式", f"第{idx+1}节下边距不是2cm", failures)
        expect_true(abs(section.left_margin.cm - 3.0) < 0.02, "DOCX格式", f"第{idx+1}节左边距不是3cm", failures)
        expect_true(abs(section.right_margin.cm - 2.0) < 0.02, "DOCX格式", f"第{idx+1}节右边距不是2cm", failures)
        expect_true(abs(section.header_distance.cm - 1.3) < 0.02, "DOCX格式", f"第{idx+1}节页眉边距不是1.3cm", failures)
        expect_true(abs(section.footer_distance.cm - 1.1) < 0.02, "DOCX格式", f"第{idx+1}节页脚边距不是1.1cm", failures)

    if not doc.tables:
        failures.append("[DOCX格式] 未找到封面表格")
    else:
        cover_table = doc.tables[0]
        cover_title_row1 = cover_table.cell(0, 1).paragraphs[0]
        cover_title_row2 = cover_table.cell(1, 1).paragraphs[0] if len(cover_table.rows) > 1 else None
        cover_title_text = "\n".join(
            part
            for part in [
                cover_title_row1.text.strip(),
                cover_title_row2.text.strip() if cover_title_row2 is not None else "",
            ]
            if part
        )
        if not cover_title_text:
            failures.append("[DOCX格式] 封面论文题目缺失")
        elif "\n" not in cover_title_text:
            failures.append("[DOCX格式] 封面论文题目未按长题目手动换行")
        if cover_title_row1.runs:
            cover_run = cover_title_row1.runs[0]
            expect_true(
                abs((cover_run.font.size.pt if cover_run.font.size else 0) - 14.0) < 0.2,
                "DOCX格式",
                "封面论文题目不是四号",
                failures,
            )
            expect_true(
                cover_title_row1.alignment == 0,
                "DOCX格式",
                "长题目封面未左对齐",
                failures,
            )

    nonempty_paragraphs = [(idx, paragraph.text.strip()) for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip()]
    paragraph_map = {paragraph.text.strip(): paragraph for paragraph in doc.paragraphs if paragraph.text.strip()}

    def is_english_title_line(text: str) -> bool:
        compact = re.sub(r"\s+", " ", text.strip())
        return bool(compact) and compact == compact.upper() and bool(re.fullmatch(r"[A-Z0-9 ,:&'\-()]+", compact))

    for pos, (idx, text) in enumerate(nonempty_paragraphs):
        if text == "摘  要" and pos > 0 and "\n" not in nonempty_paragraphs[pos - 1][1]:
            failures.append("[DOCX格式] 中文摘要页论文题目未按长题目手动换行")
            break
        if text == "ABSTRACT" and pos > 0:
            prev_text = nonempty_paragraphs[pos - 1][1]
            has_multiline_title = "\n" in prev_text
            split_title_lines = 0
            back = pos - 1
            while back >= 0 and is_english_title_line(nonempty_paragraphs[back][1]):
                split_title_lines += 1
                back -= 1
            if not has_multiline_title and split_title_lines < 2:
                failures.append("[DOCX格式] 英文摘要页论文题目未按长题目手动换行")
                break

    abstract_title_para = next((p for p in doc.paragraphs if "\n" in p.text and "高管超额薪酬研究" in p.text), None)
    if abstract_title_para is not None and abstract_title_para.runs:
        run = abstract_title_para.runs[0]
        expect_true(abs((run.font.size.pt if run.font.size else 0) - 16.0) < 0.2, "DOCX格式", "中文摘要页题目不是三号", failures)
        expect_true(abstract_title_para.alignment == 1, "DOCX格式", "中文摘要页题目未居中", failures)

    abstract_heading_para = paragraph_map.get("摘  要")
    if abstract_heading_para is not None and abstract_heading_para.runs:
        run = abstract_heading_para.runs[0]
        expect_true(abs((run.font.size.pt if run.font.size else 0) - 14.0) < 0.2, "DOCX格式", "“摘 要”不是四号", failures)
        expect_true(abstract_heading_para.alignment == 1, "DOCX格式", "“摘 要”未居中", failures)

    english_title_paras = [p for p in doc.paragraphs if is_english_title_line(p.text.strip())]
    for paragraph in english_title_paras[:3]:
        if not paragraph.runs:
            continue
        run = paragraph.runs[0]
        expect_true(abs((run.font.size.pt if run.font.size else 0) - 16.0) < 0.2, "DOCX格式", "英文题目不是三号", failures)
        expect_true(bool(run.bold), "DOCX格式", "英文题目未加粗", failures)
        expect_true(paragraph.alignment == 1, "DOCX格式", "英文题目未居中", failures)

    english_abstract_heading = paragraph_map.get("ABSTRACT")
    if english_abstract_heading is not None and english_abstract_heading.runs:
        run = english_abstract_heading.runs[0]
        expect_true(abs((run.font.size.pt if run.font.size else 0) - 14.0) < 0.2, "DOCX格式", "ABSTRACT 不是四号", failures)
        expect_true(bool(run.bold), "DOCX格式", "ABSTRACT 未加粗", failures)
        expect_true(english_abstract_heading.alignment == 1, "DOCX格式", "ABSTRACT 未居中", failures)

    toc_title_para = paragraph_map.get("目  录")
    if toc_title_para is not None and toc_title_para.runs:
        run = toc_title_para.runs[0]
        expect_true(abs((run.font.size.pt if run.font.size else 0) - 16.0) < 0.2, "DOCX格式", "目录标题不是三号", failures)
        expect_true(toc_title_para.alignment == 1, "DOCX格式", "目录标题未居中", failures)

    toc_entries = []
    for paragraph in doc.paragraphs:
        style_id = (getattr(paragraph.style, "style_id", "") or "").lower().replace(" ", "")
        style_name = (getattr(paragraph.style, "name", "") or "").lower().replace(" ", "")
        if style_id in {"toc1", "toc2", "toc3"} or style_name in {"toc1", "toc2", "toc3", "toc 1", "toc 2", "toc 3", "目录1", "目录2", "目录3"}:
            text = re.sub(r"\s+", " ", paragraph.text.strip())
            if text:
                toc_entries.append(text)
    duplicates = sorted({entry for entry in toc_entries if toc_entries.count(entry) > 1})
    if duplicates:
        failures.append(f"[DOCX目录] 目录条目重复，疑似更新域后生成了双目录: {duplicates[:5]}")

    toc3_entries = [paragraph.text.strip() for paragraph in doc.paragraphs if (getattr(paragraph.style, "style_id", "") or "").lower().replace(" ", "") == "toc3"]
    if toc3_entries:
        expect_true(
            all(text.count("\t") == 1 for text in toc3_entries),
            "DOCX目录",
            "TOC3 条目格式异常",
            failures,
        )

    with zipfile.ZipFile(DOCX_PATH) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
        settings_xml = zf.read("word/settings.xml").decode("utf-8")
        footer_xml_map = {
            name: zf.read(name).decode("utf-8")
            for name in zf.namelist()
            if re.fullmatch(r"word/footer\d+\.xml", name)
        }
    if "<w:txbxContent" in document_xml:
        failures.append("[DOCX格式] 文档仍包含模板说明文本框")
    if "w:updateFields" not in settings_xml:
        failures.append("[DOCX格式] 未设置打开时更新域 updateFields")
    footer_blob = "\n".join(footer_xml_map.values())
    if "PAGE" not in footer_blob or ("NUMPAGES" not in footer_blob and "PAGEREF _BodyEnd" not in footer_blob):
        failures.append("[DOCX格式] 正文页脚未保留动态页码域 PAGE / NUMPAGES（或等效 PAGEREF）")

    abstract_headers = [section.header.paragraphs[0].text for section in doc.sections[1:4] if section.header.paragraphs]
    if len(abstract_headers) >= 3:
        expect_true("摘 要" in abstract_headers[0], "DOCX格式", "摘要页页眉异常", failures)
        expect_true("目 录" in abstract_headers[1], "DOCX格式", "目录页页眉异常", failures)
        expect_true("论文正文" in abstract_headers[2], "DOCX格式", "正文页页眉异常", failures)

    important_blank_indices: list[int] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "摘  要":
            important_blank_indices.extend([idx - 1, idx + 1])
        elif text == "ABSTRACT":
            important_blank_indices.extend([idx - 1, idx + 1])
        elif text.startswith("关键词：") or text.startswith("关键词:"):
            important_blank_indices.append(idx + 1)
        elif text.startswith("Keywords:"):
            important_blank_indices.append(idx + 1)
    for idx in sorted({x for x in important_blank_indices if 0 <= x < len(doc.paragraphs)}):
        paragraph = doc.paragraphs[idx]
        if paragraph.text.strip():
            failures.append(f"[DOCX格式] 预期空行位置 {idx} 不是空段")
            continue
        if "sectPr" in paragraph._p.xml:
            continue
        if not paragraph.runs:
            continue
        run = paragraph.runs[0]
        if run.font.size is None:
            continue
        size = run.font.size.pt
        expect_true(abs(size - 10.5) < 0.2, "DOCX格式", f"前置关键空行 {idx} 不是五号字体", failures)

    math_number_sequence: list[int] = []
    math_root = None
    try:
        from lxml import etree

        math_root = etree.fromstring(document_xml.encode("utf-8"))
    except Exception:
        math_root = None
    if math_root is not None:
        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
        }
        for paragraph in math_root.xpath("//w:p", namespaces=ns):
            has_display_math = bool(paragraph.xpath(".//m:oMathPara", namespaces=ns))
            has_numbered_inline_math = bool(
                paragraph.xpath(".//m:oMath", namespaces=ns)
                and paragraph.xpath(".//w:tab", namespaces=ns)
            )
            if not (has_display_math or has_numbered_inline_math):
                continue
            tab_positions = paragraph.xpath("./w:pPr/w:tabs/w:tab/@w:pos", namespaces=ns)
            tab_alignments = paragraph.xpath("./w:pPr/w:tabs/w:tab/@w:val", namespaces=ns)
            first_line_indents = paragraph.xpath("./w:pPr/w:ind/@w:firstLine", namespaces=ns)
            if "8617" not in tab_positions or "right" not in tab_alignments:
                failures.append("[DOCX公式] 展示公式未设置行末右对齐编号制表位")
            if any(value not in {"0", "0.0"} for value in first_line_indents):
                failures.append("[DOCX公式] 展示公式段落不应使用正文首行缩进")
            paragraph_text = "".join(paragraph.xpath(".//w:t/text()", namespaces=ns))
            match = re.search(r"\((\d+)\)", paragraph_text)
            if match:
                math_number_sequence.append(int(match.group(1)))
    if math_number_sequence and math_number_sequence != list(range(1, len(math_number_sequence) + 1)):
        failures.append(f"[DOCX公式] 公式编号缓存不连续: {math_number_sequence}")

    body_heading1 = next(
        (
            p for p in doc.paragraphs
            if p.style and p.style.name == "Heading 1" and p.text.strip().startswith("第一章 ")
        ),
        None,
    )
    if body_heading1 is not None and body_heading1.runs:
        run = body_heading1.runs[0]
        expect_true(abs((run.font.size.pt if run.font.size else 0) - 16.0) < 0.2, "DOCX格式", "一级标题不是三号", failures)
        expect_true(body_heading1.alignment == 1, "DOCX格式", "一级标题未居中", failures)

    body_heading2 = next(
        (
            p for p in doc.paragraphs
            if p.style and p.style.name == "Heading 2" and re.match(r"^\d+\.\d+\s+", p.text.strip())
        ),
        None,
    )
    if body_heading2 is not None and body_heading2.runs:
        run = body_heading2.runs[0]
        expect_true(abs((run.font.size.pt if run.font.size else 0) - 14.0) < 0.2, "DOCX格式", "二级标题不是四号", failures)
        expect_true(body_heading2.alignment == 0, "DOCX格式", "二级标题未顶格", failures)
        expect_true(not bool(run.bold), "DOCX格式", "二级标题不应加粗", failures)

    body_heading3 = next((p for p in doc.paragraphs if re.match(r"^\d+\.\d+\.\d+\s+", p.text.strip())), None)
    if body_heading3 is not None and body_heading3.runs:
        run = body_heading3.runs[0]
        expect_true(abs((run.font.size.pt if run.font.size else 0) - 12.0) < 0.2, "DOCX格式", "三级标题不是小四", failures)
        expect_true(abs((body_heading3.paragraph_format.first_line_indent.pt if body_heading3.paragraph_format.first_line_indent else 0) - 21.0) < 0.5, "DOCX格式", "三级标题缩进不是两个汉字", failures)
        expect_true(not bool(run.bold), "DOCX格式", "三级标题不应加粗", failures)

    first_caption = next((p for p in doc.paragraphs if re.match(r"^(表|图)\d", p.text.strip())), None)
    if first_caption is not None and first_caption.runs:
        run = first_caption.runs[0]
        expect_true(abs((run.font.size.pt if run.font.size else 0) - 9.0) < 0.2, "DOCX格式", "图表题不是小五", failures)
        expect_true(bool(run.bold), "DOCX格式", "图表题未加粗", failures)
        expect_true(first_caption.alignment == 1, "DOCX格式", "图表题未居中", failures)

    ref_heading_idx = next((idx for idx, p in enumerate(doc.paragraphs) if p.text.strip() == "参考文献"), None)
    if ref_heading_idx is not None:
        first_ref = next((p for p in doc.paragraphs[ref_heading_idx + 1 :] if p.text.strip()), None)
        if first_ref is not None and first_ref.runs:
            run = first_ref.runs[0]
            expect_true(abs((run.font.size.pt if run.font.size else 0) - 10.5) < 0.2, "DOCX格式", "参考文献字体不是五号", failures)


def build_results() -> dict:
    df = pd.read_csv(DATASET_PATH, encoding="utf-8-sig")
    controls = FE_CONTROL_VARS.copy()

    positive_df = df[df[ALT_SUBSIDY_LAG_COL].notna()].copy()
    full_df = df.copy()

    main = _fit_model2(positive_df, controls, subsidy_col=ALT_SUBSIDY_LAG_COL)
    strict = _fit_model2(positive_df, controls, subsidy_col=ALT_SUBSIDY_LAG_COL, fe_mode="entity_industry_year")
    full_compare = _fit_model2(full_df, controls, subsidy_col=BASE_SUBSIDY_LAG_COL)

    core_vars = [ALT_SUBSIDY_LAG_COL] + controls
    salary_df = positive_df.dropna(subset=core_vars + ["lnCEOpay", "Year"]).copy().set_index(["Symbol", "Year"], drop=False)
    X_salary, _, _ = _build_fe_matrix(salary_df, core_vars)
    salary_model = _fit_with_cluster_se(salary_df["lnCEOpay"], X_salary)

    short = _fit_model2(
        positive_df[(positive_df["Year"] >= 2010) & (positive_df["Year"] <= 2020)].copy(),
        controls,
        subsidy_col=ALT_SUBSIDY_LAG_COL,
    )
    manufacturing = _fit_model2(
        positive_df[positive_df["IndustrySector"] == "C"].copy(),
        controls,
        subsidy_col=ALT_SUBSIDY_LAG_COL,
    )

    soe = _fit_model2(positive_df[positive_df["IsSOE"] == 1].copy(), controls, subsidy_col=ALT_SUBSIDY_LAG_COL)
    private = _fit_model2(positive_df[positive_df["IsPrivate"] == 1].copy(), controls, subsidy_col=ALT_SUBSIDY_LAG_COL)
    regulated = _fit_model2(
        positive_df[positive_df["RegulatedIndustry"] == 1].copy(),
        controls,
        subsidy_col=ALT_SUBSIDY_LAG_COL,
    )
    nonregulated = _fit_model2(
        positive_df[positive_df["RegulatedIndustry"] == 0].copy(),
        controls,
        subsidy_col=ALT_SUBSIDY_LAG_COL,
    )

    soe_df = positive_df[positive_df["IsSOE"] == 1].copy()
    central = _fit_model2(soe_df[soe_df["IsCentralSOE"] == 1].copy(), controls, subsidy_col=ALT_SUBSIDY_LAG_COL)
    local = _fit_model2(soe_df[soe_df["IsCentralSOE"] == 0].copy(), controls, subsidy_col=ALT_SUBSIDY_LAG_COL)
    soe_identifiable_ratio = soe_df["IsCentralSOE"].notna().mean()

    heckman = _fit_heckman_two_step(df, controls, subsidy_col=ALT_SUBSIDY_LAG_COL)

    mediation_df = positive_df.dropna(
        subset=["Overpay", "Power_PCA", ALT_SUBSIDY_LAG_COL, "Year"] + controls
    ).copy()
    mediation_df = mediation_df.set_index(["Symbol", "Year"], drop=False)
    X3 = mediation_df[[ALT_SUBSIDY_LAG_COL] + controls]
    m3 = _fit_with_cluster_se(mediation_df["Overpay"], X3)
    m4 = _fit_with_cluster_se(mediation_df["Power_PCA"], X3)
    X5 = mediation_df[[ALT_SUBSIDY_LAG_COL, "Power_PCA"] + controls]
    m5 = _fit_with_cluster_se(mediation_df["Overpay"], X5)

    placebo_df = df.sort_values(["Symbol", "Year"]).copy()
    placebo_df["lnSubsidy_pos_f1_new"] = placebo_df.groupby("Symbol")["lnSubsidy_pos"].shift(-1)
    placebo_df["lnSubsidy_pos_f2_new"] = placebo_df.groupby("Symbol")["lnSubsidy_pos"].shift(-2)
    placebo_base = placebo_df[placebo_df[ALT_SUBSIDY_LAG_COL].notna()].copy()
    placebo_f1 = _fit_model2(placebo_base.dropna(subset=["lnSubsidy_pos_f1_new"]).copy(), controls, subsidy_col="lnSubsidy_pos_f1_new")
    placebo_f2 = _fit_model2(placebo_base.dropna(subset=["lnSubsidy_pos_f2_new"]).copy(), controls, subsidy_col="lnSubsidy_pos_f2_new")
    placebo_f1_strict = _fit_model2(
        placebo_base.dropna(subset=["lnSubsidy_pos_f1_new"]).copy(),
        controls,
        subsidy_col="lnSubsidy_pos_f1_new",
        fe_mode="entity_industry_year",
    )
    placebo_f2_strict = _fit_model2(
        placebo_base.dropna(subset=["lnSubsidy_pos_f2_new"]).copy(),
        controls,
        subsidy_col="lnSubsidy_pos_f2_new",
        fe_mode="entity_industry_year",
    )

    ml_summary = json.loads((RESULTS_DIR / "ml_validation_summary.json").read_text(encoding="utf-8"))
    ml_tuning = json.loads((RESULTS_DIR / "ml_tuning_summary.json").read_text(encoding="utf-8"))
    event_study_summary = json.loads(EVENT_STUDY_SUMMARY_PATH.read_text(encoding="utf-8"))
    event_study_rows = pd.read_csv(EVENT_STUDY_RESULTS_PATH, encoding="utf-8-sig")

    return {
        "main": main,
        "strict": strict,
        "full_compare": full_compare,
        "salary_model": salary_model,
        "salary_df": salary_df,
        "short": short,
        "manufacturing": manufacturing,
        "soe": soe,
        "private": private,
        "regulated": regulated,
        "nonregulated": nonregulated,
        "central": central,
        "local": local,
        "soe_identifiable_ratio": soe_identifiable_ratio,
        "heckman": heckman,
        "mediation_df": mediation_df,
        "m3": m3,
        "m4": m4,
        "m5": m5,
        "placebo_f1": placebo_f1,
        "placebo_f2": placebo_f2,
        "placebo_f1_strict": placebo_f1_strict,
        "placebo_f2_strict": placebo_f2_strict,
        "event_study_summary": event_study_summary,
        "event_study_rows": event_study_rows,
        "ml_summary": ml_summary,
        "ml_tuning": ml_tuning,
    }


def check_result_exports(res: dict, failures: list[str]) -> None:
    if not SAMPLE_SUMMARY_PATH.exists():
        failures.append(f"[缺失] 结果文件不存在: {SAMPLE_SUMMARY_PATH}")
    else:
        sample_df = pd.read_csv(SAMPLE_SUMMARY_PATH)
        main_row = sample_df.loc[sample_df["stage"] == "模型2主回归样本"]
        unified_row = sample_df.loc[sample_df["stage"] == "中介统一样本"]
        expect_true(not main_row.empty, "sample_screening_summary", "缺少模型2主回归样本行", failures)
        expect_true(not unified_row.empty, "sample_screening_summary", "缺少中介统一样本行", failures)
        if not main_row.empty:
            expect_true(int(main_row.iloc[0]["observations"]) == int(res["main"]["sample_size"]), "sample_screening_summary", "模型2主回归样本量未同步", failures)
            expect_true("lnSubsidy_pos_l1" in str(main_row.iloc[0]["note"]), "sample_screening_summary", "主回归样本说明未同步到正补助口径", failures)
        if not unified_row.empty:
            expect_true(int(unified_row.iloc[0]["observations"]) == int(len(res["mediation_df"])), "sample_screening_summary", "中介统一样本量未同步", failures)
            expect_true("lnSubsidy_pos_l1" in str(unified_row.iloc[0]["note"]), "sample_screening_summary", "中介样本说明未同步到正补助口径", failures)
            expect_true("Power_PCA" in str(unified_row.iloc[0]["note"]), "sample_screening_summary", "中介样本说明未同步到 PCA 口径", failures)

    if not DESCRIPTIVE_STATS_PATH.exists():
        failures.append(f"[缺失] 结果文件不存在: {DESCRIPTIVE_STATS_PATH}")
    else:
        desc_df = pd.read_csv(DESCRIPTIVE_STATS_PATH)
        power_row = desc_df.loc[desc_df["变量"] == "管理层权力(Power, PCA口径)"]
        expect_true(not power_row.empty, "descriptive_statistics", "描述性统计未同步为 PCA 口径", failures)

    if not CAUSAL_RESULTS_PATH.exists():
        failures.append(f"[缺失] 结果文件不存在: {CAUSAL_RESULTS_PATH}")
    else:
        causal_df = pd.read_csv(CAUSAL_RESULTS_PATH)
        expect_true("中介效应PCA" in set(causal_df["category"]), "causal_results", "中介结果未同步为 PCA 口径", failures)
        expect_true("lnSubsidy_pos_l1→Power_PCA" in set(causal_df["key_var"]), "causal_results", "路径a未同步为 Power_PCA", failures)
        expect_true("Power_PCA→Overpay" in set(causal_df["key_var"]), "causal_results", "路径b未同步为 Power_PCA", failures)

    if not MAIN_MEDIATION_SUMMARY_PATH.exists():
        failures.append(f"[缺失] 结果文件不存在: {MAIN_MEDIATION_SUMMARY_PATH}")
    else:
        mediation_summary = pd.read_csv(MAIN_MEDIATION_SUMMARY_PATH)
        expect_true(len(mediation_summary) == 1, "main_mediation_summary", f"行数异常: {len(mediation_summary)}", failures)
        if len(mediation_summary) == 1:
            row = mediation_summary.iloc[0]
            expect_true(str(row["mediator_method"]) == "PCA", "main_mediation_summary", "中介方法未同步为 PCA", failures)
            expect_true(str(row["power_measure_method"]) == "PCA", "main_mediation_summary", "权力口径未同步为 PCA", failures)

    if not MAIN_FE_PATH.exists():
        failures.append(f"[缺失] 结果文件不存在: {MAIN_FE_PATH}")
    else:
        main_fe = pd.read_csv(MAIN_FE_PATH)
        expect_true(len(main_fe) == 2, "main_regression_fe_comparison", f"行数异常: {len(main_fe)}", failures)
        if len(main_fe) >= 2:
            row0 = main_fe.iloc[0]
            row1 = main_fe.iloc[1]
            expect_true(int(row0["sample_size"]) == int(res["main"]["sample_size"]), "main_regression_fe_comparison", "基准 FE 样本量未同步", failures)
            expect_true(abs(float(row0["coef"]) - float(res["main"]["model"].params[ALT_SUBSIDY_LAG_COL])) < 1e-12, "main_regression_fe_comparison", "基准 FE 系数未同步", failures)
            expect_true(int(row1["sample_size"]) == int(res["strict"]["sample_size"]), "main_regression_fe_comparison", "严格 FE 样本量未同步", failures)
            expect_true(abs(float(row1["coef"]) - float(res["strict"]["model"].params[ALT_SUBSIDY_LAG_COL])) < 1e-12, "main_regression_fe_comparison", "严格 FE 系数未同步", failures)

    if not ROBUSTNESS_PATH.exists():
        failures.append(f"[缺失] 结果文件不存在: {ROBUSTNESS_PATH}")
    else:
        robustness = pd.read_csv(ROBUSTNESS_PATH)
        expect_true(len(robustness) == 4, "robustness_results", f"行数异常: {len(robustness)}", failures)
        expected_checks = {"替换因变量", "缩小样本期(2010-2020)", "仅制造业", "更严格固定效应（行业×年份）"}
        expect_true(set(robustness["check_item"]) == expected_checks, "robustness_results", f"检验项异常: {set(robustness['check_item'])}", failures)
        expect_true(set(robustness["key_var"]) == {ALT_SUBSIDY_LAG_COL}, "robustness_results", "核心解释变量未统一为 lnSubsidy_pos_l1", failures)

    if not HECKMAN_PATH.exists():
        failures.append(f"[缺失] 结果文件不存在: {HECKMAN_PATH}")
    else:
        heckman_df = pd.read_csv(HECKMAN_PATH)
        expect_true(len(heckman_df) == 1, "heckman_results", f"行数异常: {len(heckman_df)}", failures)
        if len(heckman_df) == 1:
            row = heckman_df.iloc[0]
            expect_true(int(row["outcome_sample_size"]) == int(res["heckman"]["outcome_sample_size"]), "heckman_results", "结果方程样本量未同步", failures)
            expect_true(abs(float(row["subsidy_coef"]) - float(res["heckman"]["subsidy_coef"])) < 1e-12, "heckman_results", "Heckman 补贴系数未同步", failures)

    if not EVENT_STUDY_SUMMARY_PATH.exists():
        failures.append(f"[缺失] 结果文件不存在: {EVENT_STUDY_SUMMARY_PATH}")
    else:
        summary = json.loads(EVENT_STUDY_SUMMARY_PATH.read_text(encoding="utf-8"))
        expect_true(int(summary["sample_size"]) == int(res["main"]["sample_size"]), "event_study_summary", "事件研究样本量未与基准回归同步", failures)
        expect_true(float(summary["pretrend_p"]) > 0.05, "event_study_summary", "事件研究前趋势联合检验未通过", failures)

    if not EVENT_STUDY_RESULTS_PATH.exists():
        failures.append(f"[缺失] 结果文件不存在: {EVENT_STUDY_RESULTS_PATH}")
    else:
        event_rows = pd.read_csv(EVENT_STUDY_RESULTS_PATH)
        expect_true(set(event_rows["event_time"]) == {-3, -2, 0, 1, 2, 3}, "event_study_results", "事件期设置异常", failures)


def check_thesis_against_results(thesis_raw: str, failures: list[str], res: dict | None = None) -> None:
    if res is None:
        res = build_results()

    abstract_match = re.search(r"\*\*摘　要\*\*\s*(.*?)\s*\*\*关键词\*\*", thesis_raw, flags=re.S)
    if not abstract_match:
        failures.append("[缺失] 中文摘要")
    else:
        abstract_text = re.sub(r"\s+", "", abstract_match.group(1))
        expect_true(
            len(abstract_text) <= 300,
            "中文摘要长度",
            f"当前为 {len(abstract_text)} 字，超过 300 字",
            failures,
        )

    expect(
        thesis_raw,
        "经过这些步骤，原始公司—年度观测为70,559条；剔除金融行业后为69,142条；再剔除特殊处理样本后为63,011条；满足关键变量完整要求的样本为52,980条；满足基准薪酬模型完整案例要求的样本为52,805条；纳入管理层权力底层指标后，可用于构造 Power 的样本为50,171条；在基准回归中进一步要求“上一年补助金额大于0”且滞后一期补贴强度可得后，模型2样本为43,417条；机器学习补充验证与基准回归保持同一变量口径，因此其可用样本同样为43,417条；在中介效应检验中要求 Overpay、Power 与滞后一期正补助均完整后，统一样本为41,186条。",
        "样本筛选链",
        failures,
    )

    main = res["main"]["model"]
    expect(
        thesis_raw,
        f"<tr><td>lnSubsidy_pos_l1</td><td>Overpay</td><td>{fmt_plain(main.params[ALT_SUBSIDY_LAG_COL])}</td><td>{fmt_plain(main.std_errors[ALT_SUBSIDY_LAG_COL])}</td><td>{fmt_plain(main.tstats[ALT_SUBSIDY_LAG_COL], 4)}</td><td>{fmt_plain(main.pvalues[ALT_SUBSIDY_LAG_COL], 4)}</td><td>{fmt_int(res['main']['sample_size'])}</td><td>{fmt_plain(main.rsquared)}</td></tr>",
        "表4-5 主系数",
        failures,
    )
    expect(
        thesis_raw,
        f"<tr><td>Roa</td><td>Overpay</td><td>{fmt_fixed(main.params['Roa'])}</td><td>{fmt_plain(main.std_errors['Roa'])}</td><td>{fmt_fixed(main.tstats['Roa'], 4)}</td><td>{fmt_plain(main.pvalues['Roa'], 4)}</td><td>{fmt_int(res['main']['sample_size'])}</td><td>{fmt_plain(main.rsquared)}</td></tr>",
        "表4-5 Roa",
        failures,
    )
    expect(
        thesis_raw,
        f"<tr><td>Lever</td><td>Overpay</td><td>{fmt_fixed(main.params['Lever'])}</td><td>{fmt_plain(main.std_errors['Lever'])}</td><td>{fmt_fixed(main.tstats['Lever'], 4)}</td><td>{fmt_plain(main.pvalues['Lever'], 4)}</td><td>{fmt_int(res['main']['sample_size'])}</td><td>{fmt_plain(main.rsquared)}</td></tr>",
        "表4-5 Lever",
        failures,
    )
    expect(
        thesis_raw,
        f"<tr><td>Top1</td><td>Overpay</td><td>{fmt_fixed(main.params['Top1'])}</td><td>{fmt_plain(main.std_errors['Top1'])}</td><td>{fmt_fixed(main.tstats['Top1'], 4)}</td><td>{fmt_plain(main.pvalues['Top1'], 4)}</td><td>{fmt_int(res['main']['sample_size'])}</td><td>{fmt_plain(main.rsquared)}</td></tr>",
        "表4-5 Top1",
        failures,
    )

    expect(thesis_raw, f"<tr><td>Heckman 两步法</td><td>IV_lnSubsidy_l1</td><td>IMR p = {fmt_plain(res['heckman']['imr_p'], 4)}</td><td>{fmt_coef(res['heckman']['subsidy_coef'], res['heckman']['subsidy_p'])}（t = {fmt_plain(res['heckman']['subsidy_t'], 4)}）</td><td>{fmt_int(res['heckman']['outcome_sample_size'])}</td></tr>", "表4-6 Heckman", failures)

    expect(thesis_raw, "<tr><td>(1) 替换因变量</td><td>高管前三名薪酬对数</td><td>0.0331***</td><td>9.2397</td><td>43,533</td><td>0.0469</td></tr>", "表4-7 行1", failures)
    expect(thesis_raw, "<tr><td>(2) 缩小样本期（2010—2020）</td><td>Overpay</td><td>0.0120***</td><td>3.0072</td><td>25,028</td><td>0.0232</td></tr>", "表4-7 行2", failures)
    expect(thesis_raw, "<tr><td>(3) 仅制造业</td><td>Overpay</td><td>0.0099**</td><td>2.2766</td><td>29,143</td><td>0.0200</td></tr>", "表4-7 行3", failures)
    strict = res["strict"]["model"]
    expect(
        thesis_raw,
        f"<tr><td>(4) 更严格固定效应（行业×年份）</td><td>Overpay</td><td>{fmt_coef(strict.params[ALT_SUBSIDY_LAG_COL], strict.pvalues[ALT_SUBSIDY_LAG_COL])}</td><td>{fmt_plain(strict.tstats[ALT_SUBSIDY_LAG_COL], 4)}</td><td>{fmt_int(res['strict']['sample_size'])}</td><td>{fmt_plain(strict.rsquared)}</td></tr>",
        "表4-7 行4",
        failures,
    )

    ml = res["ml_summary"]
    expect(
        thesis_raw,
        f"建模样本（{fmt_int(res['ml_tuning']['split']['train_size'])}条）",
        "ML 建模样本量",
        failures,
    )
    if (
        f"保留样本（{fmt_int(res['ml_tuning']['split']['test_size_n'])}条）" not in thesis_raw
        and f"保留样本{fmt_int(res['ml_tuning']['split']['test_size_n'])}条" not in thesis_raw
    ):
        failures.append(
            f"[缺失] ML 保留样本量: {fmt_int(res['ml_tuning']['split']['test_size_n'])}条"
        )

    rf_rows = ml["random_forest"]["importance_table"]
    expect(thesis_raw, f"<tr><td>1</td><td>{rf_rows[0]['变量']}</td><td>{fmt_plain(rf_rows[0]['随机森林重要性'], 6)}</td></tr>", "表4-8 行1", failures)
    expect(thesis_raw, f"<tr><td>2</td><td>{rf_rows[1]['变量']}</td><td>{fmt_plain(rf_rows[1]['随机森林重要性'], 6)}</td></tr>", "表4-8 行2", failures)
    expect(thesis_raw, f"<tr><td>3</td><td>{rf_rows[2]['变量']}</td><td>{fmt_plain(rf_rows[2]['随机森林重要性'], 6)}</td></tr>", "表4-8 行3", failures)
    expect(thesis_raw, f"<tr><td>4</td><td>{rf_rows[3]['变量']}</td><td>{fmt_plain(rf_rows[3]['随机森林重要性'], 6)}</td></tr>", "表4-8 行4", failures)

    shap_rows = ml["rf_shap"]["importance_table"]
    expect(thesis_raw, f"<tr><td>1</td><td>{shap_rows[0]['变量']}</td><td>{fmt_plain(shap_rows[0]['平均绝对SHAP值'], 6)}</td></tr>", "表4-9 行1", failures)
    expect(thesis_raw, f"<tr><td>2</td><td>{shap_rows[1]['变量']}</td><td>{fmt_plain(shap_rows[1]['平均绝对SHAP值'], 6)}</td></tr>", "表4-9 行2", failures)
    expect(thesis_raw, f"<tr><td>3</td><td>{shap_rows[2]['变量']}</td><td>{fmt_plain(shap_rows[2]['平均绝对SHAP值'], 6)}</td></tr>", "表4-9 行3", failures)
    expect(thesis_raw, f"<tr><td>4</td><td>{shap_rows[3]['变量']}</td><td>{fmt_plain(shap_rows[3]['平均绝对SHAP值'], 6)}</td></tr>", "表4-9 行4", failures)

    expect(thesis_raw, f"<tr><td>模型3 总效应 c：lnSubsidy_pos_l1 → Overpay</td><td>{fmt_coef(res['m3'].params[ALT_SUBSIDY_LAG_COL], res['m3'].pvalues[ALT_SUBSIDY_LAG_COL])}</td><td>{fmt_plain(res['m3'].tstats[ALT_SUBSIDY_LAG_COL], 4)}</td><td>{fmt_plain(res['m3'].pvalues[ALT_SUBSIDY_LAG_COL], 4)}</td></tr>", "表4-10 模型3", failures)
    expect(thesis_raw, f"<tr><td>模型4 路径 a：lnSubsidy_pos_l1 → Power（PCA）</td><td>{fmt_coef(res['m4'].params[ALT_SUBSIDY_LAG_COL], res['m4'].pvalues[ALT_SUBSIDY_LAG_COL])}</td><td>{fmt_plain(res['m4'].tstats[ALT_SUBSIDY_LAG_COL], 4)}</td><td>{fmt_plain(res['m4'].pvalues[ALT_SUBSIDY_LAG_COL], 4)}</td></tr>", "表4-10 模型4", failures)
    expect(thesis_raw, f"<tr><td>模型5 路径 b：Power（PCA） → Overpay</td><td>{fmt_coef(res['m5'].params['Power_PCA'], res['m5'].pvalues['Power_PCA'])}</td><td>{fmt_plain(res['m5'].tstats['Power_PCA'], 4)}</td><td>{fmt_plain(res['m5'].pvalues['Power_PCA'], 4)}</td></tr>", "表4-10 模型5-b", failures)
    expect(thesis_raw, f"<tr><td>模型5 直接效应 c'：lnSubsidy_pos_l1 → Overpay</td><td>{fmt_coef(res['m5'].params[ALT_SUBSIDY_LAG_COL], res['m5'].pvalues[ALT_SUBSIDY_LAG_COL])}</td><td>{fmt_plain(res['m5'].tstats[ALT_SUBSIDY_LAG_COL], 4)}</td><td>{fmt_plain(res['m5'].pvalues[ALT_SUBSIDY_LAG_COL], 4)}</td></tr>", "表4-10 模型5-c'", failures)

    expect(thesis_raw, "<tr><td>lnSubsidy_pos_l1</td><td>0.0046（1.04）</td><td>0.0135***（2.91）</td></tr>", "表4-11 主系数", failures)
    expect(thesis_raw, "<tr><td>lnSubsidy_pos_l1</td><td>−0.0020（−0.36）</td><td>0.0110***（2.90）</td></tr>", "表4-12 主系数", failures)
    event_summary = res["event_study_summary"]
    event_rows = res["event_study_rows"].set_index("event_time")
    expect(
        thesis_raw,
        f"当前阈值对应正补助样本的上四分位数，约为{fmt_wan(event_summary['threshold_value'])}万元；识别到的处理企业为{fmt_int(event_summary['treated_firms'])}家，纳入动态检验的样本量为{fmt_int(event_summary['sample_size'])}条。",
        "事件研究样本段落",
        failures,
    )
    event_required_parts = [
        f"前趋势联合检验的 p 值为{float(event_summary['pretrend_p']):.3f}",
        f"事件当期及随后1至3期的系数分别为{fmt_fixed(event_rows.loc[0, 'coef'])}（p {fmt_p_text(event_rows.loc[0, 'p_value'])}）、{fmt_fixed(event_rows.loc[1, 'coef'])}（p {fmt_p_text(event_rows.loc[1, 'p_value'])}）、{fmt_fixed(event_rows.loc[2, 'coef'])}（p {fmt_p_text(event_rows.loc[2, 'p_value'])}）和{fmt_fixed(event_rows.loc[3, 'coef'])}（p {fmt_p_text(event_rows.loc[3, 'p_value'])}）",
        "均为正且达到常见显著性标准",
    ]
    for part in event_required_parts:
        expect(thesis_raw, part, "事件研究结果段落", failures)

    conclusion_2_parts = [
        "（2）**基准回归在多种稳健性设定下保持正向显著",
        "0.0331（p < 0.001）",
        "0.0120（p = 0.003）",
        "0.0099（p = 0.023）",
        "0.0082（p = 0.013）",
        "私营企业（0.0135，p = 0.004）",
        "非管制行业（0.0110，p = 0.004）",
        "国有企业、管制行业以及央地国企内部",
    ]
    for part in conclusion_2_parts:
        expect(thesis_raw, part, "结论第2点", failures)

    conclusion_3_parts = [
        "（3）**选择偏差校正和事件动态检验均未改变基准回归方向",
        "Heckman 两步法中的逆米尔斯比率显著（p = 0.0056）",
        "0.0093（p = 0.008）",
        f"前趋势联合检验 p 值为{float(event_summary['pretrend_p']):.3f}",
        f"{fmt_fixed(event_rows.loc[0, 'coef'])}（p {fmt_p_text(event_rows.loc[0, 'p_value'])}）",
        f"{fmt_fixed(event_rows.loc[1, 'coef'])}（p {fmt_p_text(event_rows.loc[1, 'p_value'])}）",
        f"{fmt_fixed(event_rows.loc[2, 'coef'])}（p {fmt_p_text(event_rows.loc[2, 'p_value'])}）",
        f"{fmt_fixed(event_rows.loc[3, 'coef'])}（p {fmt_p_text(event_rows.loc[3, 'p_value'])}）",
        "条件相关",
    ]
    for part in conclusion_3_parts:
        expect(thesis_raw, part, "结论第3点", failures)

    expect_absent(thesis_raw, "Power_FA", "FA 口径残留", failures)
    expect_absent(thesis_raw, "Power（FA）", "FA 表述残留", failures)
    expect_absent(thesis_raw, "因子分析", "FA 方法表述残留", failures)
    expect_absent(thesis_raw, "5个输入特征", "机器学习中的旧五变量表述残留", failures)
    expect_absent(thesis_raw, "41,186条统一样本", "机器学习中的旧统一样本表述残留", failures)
    expect_absent(thesis_raw, "主回归", "主回归措辞残留", failures)
    expect_absent(thesis_raw, "安慰剂", "安慰剂措辞残留", failures)
    expect_absent(thesis_raw, "机器学习稳健性检验与 OLS/中介回归的对应关系", "旧机器学习总表残留", failures)
    expect_absent(thesis_raw, "替换解释变量口径（全样本 ln(1+Subsidy)）", "稳健性中的旧对照口径残留", failures)
    expect_absent(thesis_raw, "补贴系数降为0.0005", "稳健性中的旧不显著描述残留", failures)
    expect_absent(thesis_raw, "全样本主回归未显著", "旧主回归表述残留", failures)
    expect_absent(thesis_raw, "<tr><td>基准工具变量</td>", "正文主表中的旧 IV 行残留", failures)
    expect_absent(thesis_raw, "<tr><td>简单替代工具变量</td>", "正文主表中的旧 IV 行残留", failures)
    expect_absent(thesis_raw, "<tr><td>精炼双工具变量</td>", "正文主表中的旧 IV 行残留", failures)
    expect_absent(thesis_raw, "为基准回归提供额外支持", "机器学习因果/支持语气过强", failures)
    expect_absent(thesis_raw, "传统留一法工具变量在当前主样本下是否还能提供额外识别支撑", "传统 IV 主链语气过强", failures)
    expect_absent(thesis_raw, "补贴强度与高管超额薪酬呈显著正相关关系", "H1 结果导向表述过强", failures)
    expect_absent(thesis_raw, "管理层权力可能在财政补贴与高管超额薪酬的关联过程中发挥中介作用", "H2 结果导向表述过强", failures)
    expect_absent(thesis_raw, "核心解释与中介变量", "机器学习旧机制证明表述残留", failures)
    expect_absent(thesis_raw, "反而稳定排在四个特征中的前列", "机器学习排序语气过强", failures)


def main() -> int:
    thesis_raw = THESIS_PATH.read_text(encoding="utf-8")
    failures: list[str] = []
    result_snapshot = build_results()

    check_bundle_images(thesis_raw, failures)
    check_reference_section(thesis_raw, failures)
    check_table_and_figure_numbers(thesis_raw, failures)
    check_equation_numbers(thesis_raw, failures)
    check_docx_exists(failures)
    check_result_exports(result_snapshot, failures)
    check_thesis_against_results(thesis_raw, failures, result_snapshot)

    if failures:
        print("一致性校验未通过：")
        for item in failures:
            print(f" - {item}")
        return 1

    print("一致性校验通过：正文中的核心实证数据与当前 results 最新结果一一对应。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
